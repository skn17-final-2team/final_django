from django.shortcuts import redirect, render
from django.views.generic import TemplateView
from apps.core.views import LoginRequiredSessionMixin
from apps.accounts.models import Dept, User
from django.db.models import Prefetch, Q

from django.shortcuts import get_object_or_404
from django.http import JsonResponse, HttpResponse

from .models import Meeting, Attendee, Domain
from django.contrib import messages
from django.db import transaction


from apps.meetings.utils.s3_upload import upload_raw_file_bytes, get_presigned_url
from apps.meetings.utils.runpod import get_stt, runpod_health

from apps.meetings.models import S3File
from django.views.decorators.http import require_GET
from datetime import date

from django.template.loader import render_to_string  
from io import BytesIO
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from django.utils.html import strip_tags
import json
import math
import re
import ast
from typing import Dict

# 한글 폰트 등록 (맑은 고딕 사용)
KOREAN_FONT_NAME = "MalgunGothic"

def _register_korean_font():
    from pathlib import Path

    # Windows 기본 맑은 고딕 경로
    default_path = Path(r"C:\Windows\Fonts\malgun.ttf")

    if default_path.exists():
        try:
            pdfmetrics.registerFont(TTFont(KOREAN_FONT_NAME, str(default_path)))
        except Exception:
            # 이미 등록되어 있거나 오류가 나도 앱 전체가 죽지 않게 함
            pass

# 모듈 import 시 한 번 호출
_register_korean_font()


# 회의 목록에서 쓸 데이터 생성하는 함수
def build_meeting_list_context(meeting_qs, login_user_id=None):
    meetings_data = []

    # 로그인 사용자 객체(필요하면)
    login_user = None
    if login_user_id:
        login_user = User.objects.select_related("dept").get(user_id=login_user_id)

    for m in meeting_qs:
        attendees = list(m.attendees.all())
        attendee_count = len(attendees)
        attendee_names = ", ".join(a.user.name for a in attendees)

        # 참여 여부 (All 페이지에서만 실제로 사용, Mine/Dept 에선 옵션)
        is_joined = False
        if login_user_id:
            if str(m.host_id) == str(login_user_id):
                is_joined = True
            else:
                is_joined = any(str(a.user_id) == str(login_user_id) for a in attendees)


        meetings_data.append(
            {
                "meeting_id": m.meeting_id,
                "title": m.title,
                "meet_date_time": m.meet_date_time,
                "place": m.place,
                "host_name": m.host.name,
                "attendee_count": attendee_count,
                "attendee_names": attendee_names,
                "is_joined": is_joined,
            }
        )

    return meetings_data, login_user

class MeetingListAllView(LoginRequiredSessionMixin, TemplateView):
    template_name = "meetings/meeting_list_all.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        login_user_id = self.request.session.get("login_user_id")

        meeting_qs = (
            Meeting.objects
            .select_related("host")
            .prefetch_related(
                Prefetch(
                    "attendees",
                    queryset=Attendee.objects.select_related("user", "user__dept"),
                )
            )
            .order_by("-meet_date_time")
        )

        meetings_data, login_user = build_meeting_list_context(
            meeting_qs, login_user_id
        )

        context["login_user"] = login_user
        context["meetings"] = meetings_data
        return context


class MeetingListMineView(LoginRequiredSessionMixin, TemplateView):
    template_name = "meetings/meeting_list_mine.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        login_user_id = self.request.session.get("login_user_id")
        if not login_user_id:
            # 이 View는 LoginRequiredMixin 이라 실제로는 여기 안 오겠지만 방어 코드
            context["login_user"] = None
            context["meetings"] = []
            return context

        meeting_qs = (
            Meeting.objects
            .select_related("host")
            .prefetch_related(
                Prefetch(
                    "attendees",
                    queryset=Attendee.objects.select_related("user", "user__dept"),
                )
            )
            .filter(
                Q(host_id=login_user_id) | Q(attendees__user_id=login_user_id)
            )
            .order_by("-meet_date_time")
            .distinct()
        )

        meetings_data, login_user = build_meeting_list_context(
            meeting_qs, login_user_id
        )

        context["login_user"] = login_user
        context["meetings"] = meetings_data
        return context

class MeetingListDeptView(LoginRequiredSessionMixin, TemplateView):
    template_name = "meetings/meeting_list_dept.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        # 1) 로그인 사용자
        login_user_id = self.request.session.get("login_user_id")
        login_user = None
        if login_user_id:
            login_user = User.objects.select_related("dept").get(user_id=login_user_id)

        # 2) 열람 가능한 회의:
        #    host도 아니고, attendee_tbl에도 없는 회의만(즉, 부서원이 참여한 회의임)
        meeting_qs = (
            Meeting.objects
            .select_related("host")
            .prefetch_related(
                Prefetch(
                    "attendees",
                    queryset=Attendee.objects.select_related("user", "user__dept"),
                )
            )
            .exclude(
                Q(host_id=login_user_id) | Q(attendees__user_id=login_user_id)
            )
            .order_by("-meet_date_time")
            .distinct()
        )

        # 3) 템플릿용 데이터 가공 (전체/내가참여한 과 동일 구조)
        meetings_data = []

        for m in meeting_qs:
            attendees = list(m.attendees.all())
            attendee_count = len(attendees)
            attendee_names = ", ".join(a.user.name for a in attendees)

            # 이 View는 "미참여 회의"라 is_joined는 항상 False
            is_joined = False

            meetings_data.append(
                {
                    "meeting_id": m.meeting_id,
                    "title": m.title,
                    "meet_date_time": m.meet_date_time,
                    "place": m.place,
                    "host_name": m.host.name,
                    "attendee_count": attendee_count,
                    "attendee_names": attendee_names,
                    "is_joined": is_joined,
                }
            )

        context["login_user"] = login_user
        context["meetings"] = meetings_data
        return context

class MeetingCreateView(LoginRequiredSessionMixin, TemplateView):
    template_name = "meetings/meeting_create.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        # DB에서 실제 부서 + 유저들을 가져오는 부분
        dept_qs = Dept.objects.prefetch_related(
            Prefetch("users", queryset=User.objects.order_by("name"))
        ).order_by("dept_name")

        context["departments"] = dept_qs   # 템플릿으로 넘김
        # 로그인한 사용자 정보도 템플릿으로 내려줌
        login_user_id = self.request.session.get("login_user_id")
        if login_user_id:
            context["login_user"] = User.objects.get(user_id=login_user_id)
        return context

    def post(self, request, *args, **kwargs):

        attendee_ids = request.POST.getlist("attendees")
        domain_names = request.POST.getlist("domains")
        if not attendee_ids:
            messages.error(request, "참석자를 최소 1명 이상 선택해 주세요.")
            context = self.get_context_data()
            context["form_title"] = request.POST.get("title", "")
            context["form_meet_date_time"] = request.POST.get("meet_date_time", "")
            context["form_place"] = request.POST.get("place", "")
            return render(request, self.template_name, context)
        
        # 1. 폼 데이터
        title = request.POST.get("title")
        meet_date_time = request.POST.get("meet_date_time")
        place = request.POST.get("place")

        # 2. 세션에서 로그인한 사용자 id 가져오기
        login_user_id = request.session.get("login_user_id")
        if not login_user_id:
            # 혹시 세션 끊긴 경우 대비
            return redirect("login")

        # 3. host로 쓸 User 객체 조회
        host_user = User.objects.get(user_id=login_user_id)
        
        attendee_ids = [uid for uid in attendee_ids if str(uid) != str(login_user_id)]


        # 4. meeting_tbl에 새 레코드 생성
        with transaction.atomic():
            # meeting_tbl insert
            meeting = Meeting.objects.create(
                title=title,
                meet_date_time=meet_date_time,
                place=place,
                host=host_user,   # FK: User 인스턴스
                transcript="",    # NOT NULL 필드라면 임시값
                domain_upload=bool(domain_names),
            )

        Attendee.objects.create(meeting=meeting, user=host_user)

        users = User.objects.filter(user_id__in=attendee_ids)
        attendee_objs = [
                Attendee(meeting=meeting, user=u) for u in users
        ]
        Attendee.objects.bulk_create(attendee_objs)

        # domain_tbl insert (특화 도메인)
        if domain_names:
            domain_objs = [
                Domain(meeting=meeting, domain_name=name)
                for name in domain_names
            ]
            Domain.objects.bulk_create(domain_objs)

        # 5. 생성된 meeting_id를 가지고 녹음 화면으로 이동
        return redirect("meetings:meeting_record", meeting_id=meeting.meeting_id)
    
class MeetingRecordView(LoginRequiredSessionMixin, TemplateView):
    template_name = "meetings/meeting_record.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        meeting_id = self.kwargs.get("meeting_id")

        meeting = Meeting.objects.get(pk=meeting_id)

        context["attendees"] = (
            meeting.attendees
                   .select_related("user")  # Attendee.user
                   .all()
        )

        context["meeting"] = meeting
        context["meeting_id"] = meeting_id
        return context

class MeetingTranscriptView(LoginRequiredSessionMixin, TemplateView):
    template_name = "meetings/meeting_transcript.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        meeting_id = self.kwargs.get("meeting_id")
        context["meeting_id"] = meeting_id
        return context

# @require_GET  
def meeting_transcript_api(request, meeting_id):
    meeting = get_object_or_404(Meeting, pk=meeting_id)
    meeting = Meeting.objects.get(pk=meeting_id)
    attendees_qs = (meeting.attendees.select_related("user", "user__dept").all())

    keys = set()
    for f in ast.literal_eval(meeting.transcript):
        for key, _ in f.items():
            keys.add(key)
    speakers = sorted(keys)
    speakers = sorted(keys)

    return JsonResponse({
        "meeting_title": meeting.title,
        "transcript": meeting.transcript,
        "record_url": meeting.record_url,
        "attendees": [{"user_id": a.user_id,"name": a.user.name, "dept_name": a.user.dept.dept_name} for a in attendees_qs],
        "speakers": speakers
    })



class MeetingDetailView(LoginRequiredSessionMixin, TemplateView):
    template_name = "meetings/meeting_detail.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        meeting_id = self.kwargs.get("meeting_id")
        meeting = get_object_or_404(Meeting, pk=meeting_id)

        # 템플릿에서 사용할 데이터 주입
        context["meeting"] = meeting
        context["attendees"] = (
            meeting.attendees
                   .select_related("user", "user__dept")
                   .all()
        )

        return context

def meeting_record_upload(request, meeting_id):
    # 1) 메소드 체크
    if request.method != "POST":
        return JsonResponse({"error": "Invalid method"}, status=405)

    # 2) 회의 조회
    meeting = get_object_or_404(Meeting, pk=meeting_id)

    # 3) 파일 추출
    uploaded_file = request.FILES.get("file")
    if not uploaded_file:
        return JsonResponse({"error": "file is required"}, status=400)

    # 4) 확장자 검증 (wav)
    filename = uploaded_file.name
    ext = filename.split(".")[-1].lower()
    if ext not in ["wav"]:
        return JsonResponse(
            {"error": "WAV 파일만 업로드 가능합니다."},
            status=400,
        )

    # 5) 유틸 호출
    file_bytes = uploaded_file.read()
    try:
        s3_key = upload_raw_file_bytes(
            file_bytes=file_bytes,
            original_filename=filename,
            # delete_after_seconds=172800, # 48시간 뒤 삭제
            delete_after_seconds=3600, # 테스트용 1시간 뒤 삭제
        )
    except Exception as e:
        # 유틸 호출 중 에러가 나도 반드시 응답을 반환
        return JsonResponse(
            {"error": f"S3 업로드 중 오류: {str(e)}"},
            status=500,
        )

    # 6) Meeting FK 연결 (record_url 이 ForeignKey(S3File, db_column="record_url") 일 때)
    meeting.record_url_id = s3_key
    meeting.save(update_fields=["record_url"])

    return JsonResponse({
        "ok": True,
    })

def meeting_summary(request, meeting_id):
    meeting = get_object_or_404(Meeting, pk=meeting_id)

    return render(request, "meetings/meeting_summary.html", {
        "meeting": meeting,
    })

class MeetingRenderingView(LoginRequiredSessionMixin, TemplateView):
    template_name = "meetings/rendering.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        meeting_id = self.kwargs.get("meeting_id")

        meeting = get_object_or_404(Meeting, pk=meeting_id)

        context["meeting"] = meeting
        context["meeting_id"] = meeting_id
        return context
    

@require_GET
def meeting_transcript_prepare(request, meeting_id):
    """
    렌딩 페이지에서 호출하는 엔드포인트.
    - 아직 transcript가 없으면 STT를 돌려서 생성
    - 이미 있으면 바로 done 리턴
    """
    meeting = get_object_or_404(Meeting, pk=meeting_id)

    transcript_html = meeting.transcript or ""

    if not transcript_html:
        # STT 호출
        print(str(meeting.record_url))
        presigned_url = get_presigned_url(str(meeting.record_url))
        print(presigned_url)
        res = get_stt(presigned_url)

        if res.status_code != 200 or not res.json().get("success"):
            # 실패 시 프론트에서 메시지 보여줄 수 있도록 에러 내려줌
            return JsonResponse(
                {
                    "status": "error",
                    "message": "전사 처리 중 오류가 발생했습니다.",
                },
                status=500,
            )
        res = res.json()
        transcript_html = res['data']['full_text']
        meeting.transcript = transcript_html
        meeting.save(update_fields=["transcript"])

    # 여기까지 왔다면 transcript 는 채워진 상태
    return JsonResponse({"status": "done"})


def today_meetings(request):
    """
    모든 템플릿에서 'today_meetings'로
    meet_date_time 기준 '오늘 날짜'인 회의 목록에 접근할 수 있게 해주는 컨텍스트 프로세서
    """
    today = date.today()

    meetings = (
        Meeting.objects
        .filter(meet_date_time__date=today)
        .select_related("host")
        .order_by("-meet_date_time")[:3]
    )

    return {
        "today_meetings": meetings
    }

def minutes_save(request, meeting_id):
    """
    회의록 HTML(meeting_notes)을 저장하는 용도
    """
    if request.method != "POST":
        return JsonResponse({"error": "Invalid method"}, status=405)

    meeting = get_object_or_404(Meeting, pk=meeting_id)

    try:
        data = json.loads(request.body.decode("utf-8"))
    except json.JSONDecodeError:
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    content = data.get("content", "")

    meeting.meeting_notes = content
    meeting.save(update_fields=["meeting_notes"])

    return JsonResponse({"ok": True})


def parse_minutes_sections(html: str) -> Dict[str, str]:
    """
    meeting.meeting_notes 에 저장된 HTML에서 data-minutes-section 별 텍스트 추출
    """
    if not html:
        return {}

    pattern = re.compile(
        r'<div[^>]*data-minutes-section="(?P<key>[^"]+)"[^>]*>(?P<body>.*?)</div>',
        re.DOTALL | re.IGNORECASE,
    )

    sections: Dict[str, str] = {}
    for match in pattern.finditer(html):
        key = match.group("key")
        body_html = match.group("body")
        text = strip_tags(body_html).replace("&nbsp;", " ")
        lines = [ln.rstrip() for ln in text.splitlines()]
        text_clean = "\n".join(ln for ln in lines if ln.strip())
        sections[key] = text_clean

    return sections

def extract_major_agenda(html: str) -> str:
    """
    회의록 HTML 안에서 '주요안건' 행 아래 minutes-editbox 내용을 우선 추출
    """
    if not html:
        return ""

    pattern = re.compile(
        r'<th[^>]*>\s*주요안건\s*</th>.*?<div[^>]*class="[^"]*minutes-editbox[^"]*"[^>]*>(?P<body>.*?)</div>',
        re.DOTALL | re.IGNORECASE,
    )
    match = pattern.search(html)
    if not match:
        return ""

    text = strip_tags(match.group("body")).replace("&nbsp;", " ").strip()
    return text


def minutes_download(request, meeting_id, fmt):
    meeting = get_object_or_404(Meeting, pk=meeting_id)

    attendees_qs = (
        meeting.attendees
               .select_related("user", "user__dept")
               .all()
    )
    attendees_count = attendees_qs.count()

    sections = parse_minutes_sections(meeting.meeting_notes or "")

    raw_contents = sections.get("contents", "")
    raw_results = sections.get("results", "")
    raw_todos = sections.get("todos", "")
    raw_base = sections.get("base", "")

    def clean_section(text: str, header_keywords) -> str:
        if not text:
            return ""
        lines = [ln.rstrip() for ln in text.splitlines()]
        while lines and not lines[0].strip():
            lines.pop(0)

        if lines:
            first = lines[0].replace(" ", "")
            if all(keyword in first for keyword in header_keywords):
                lines.pop(0)

        return "\n".join(lines).strip()

    contents_text = clean_section(raw_contents, ["회의", "내용"])
    results_text = clean_section(raw_results, ["회의", "결과"])
    todos_text = clean_section(raw_todos, ["해야", "할", "일"])
    base_text = clean_section(raw_base, ["기본", "정보"])



    has_base = "base" in sections
    has_contents = "contents" in sections
    has_results = "results" in sections
    has_todos = "todos" in sections
    has_attendees_section = "attendees" in sections

    main_agenda = extract_major_agenda(meeting.meeting_notes or "")
    if not main_agenda and base_text:
        for ln in base_text.splitlines():
            if ln.strip():
                main_agenda = ln.strip()
                break
    if not main_agenda and contents_text:
        for ln in contents_text.splitlines():
            if ln.strip():
                main_agenda = ln.strip()
                break

    agenda = meeting.title or ""
    meeting_dt_str = (
        meeting.meet_date_time.strftime("%Y.%m.%d %H:%M")
        if meeting.meet_date_time
        else ""
    )
    place = meeting.place or ""
    if hasattr(meeting, "host") and meeting.host:
        host_name = meeting.host.name
    else:
        host_name = getattr(meeting, "responsible_name", "")

    if meeting.meeting_notes:
        has_full_minutes = True
        raw_html = (
            meeting.meeting_notes
            .replace("<br>", "\n")
            .replace("<br/>", "\n")
            .replace("<br />", "\n")
        )
        text_body = strip_tags(raw_html).replace("&nbsp;", " ")
    else:
        has_full_minutes = False
        if meeting.summary:
            text_body = meeting.summary
        elif meeting.transcript:
            text_body = meeting.transcript
        else:
            text_body = "회의 내용이 아직 등록되지 않았습니다."

    if fmt == "pdf":
        buffer = BytesIO()
        p = canvas.Canvas(buffer, pagesize=A4)

        width, height = A4
        x_margin = 50
        top_margin = height - 50
        bottom_margin = 40
        line_height = 14

        font = KOREAN_FONT_NAME

        def draw_page_border(top_y, bottom_y):
            p.line(x_left, top_y, x_left, bottom_y)
            p.line(x_right, top_y, x_right, bottom_y)

        def start_new_page():
            nonlocal current_top, table_top_y
            p.setFont(font, 11)
            current_top = top_margin
            table_top_y = current_top

        def finalize_current_page(outer_bottom_value):
            draw_page_border(table_top_y, outer_bottom_value)

        y = top_margin
        p.setFont(font, 18)
        p.drawString(x_margin, y, meeting.title or "회의록")
        y -= 40

        p.setFont(font, 11)
        x_left = x_margin
        x_right = width - x_margin

        header_row_h = 24
        label_row_h = 24
        contents_h = 200
        results_h = 200
        todos_h = 150
        attend_header_h = 24
        attend_row_h = 24

        rows_per_side = max(4, math.ceil(attendees_count / 2))
        attend_rows_h = rows_per_side * attend_row_h

        table_top_y = y
        current_top = table_top_y

        center_x = (x_left + x_right) / 2.0

        if has_base:
            label_w = 80
            right_w = 120
            x_label = x_left + label_w
            x_right_block = x_right - right_w

            header_top = current_top

            for i in range(5):
                y_line = header_top - header_row_h * i
                p.line(x_left, y_line, x_right, y_line)

            p.line(x_left, header_top, x_left, header_top - 4 * header_row_h)
            p.line(x_right, header_top, x_right, header_top - 4 * header_row_h)

            p.line(x_label, header_top, x_label, header_top - header_row_h)
            p.line(x_label, header_top - header_row_h, x_label, header_top - 2 * header_row_h)
            p.line(x_label, header_top - 2 * header_row_h, x_label, header_top - 3 * header_row_h)
            p.line(x_label, header_top - 3 * header_row_h, x_label, header_top - 4 * header_row_h)

            p.line(x_right_block, header_top - header_row_h, x_right_block, header_top - 3 * header_row_h)

            def header_text_y(row_idx: int) -> float:
                return header_top - header_row_h * row_idx - header_row_h + 8

            y_row1 = header_text_y(0)
            p.drawString(x_left + 5, y_row1, "안 건")
            p.drawString(x_left + 5 + 80, y_row1, agenda)

            y_row2 = header_text_y(1)
            p.drawString(x_left + 5, y_row2, "일 시")
            p.drawString(x_left + 5 + 80, y_row2, meeting_dt_str)
            host_label = "주최자명"
            if host_name:
                host_label += f" {host_name}"
            p.drawString(x_right_block + 5, y_row2, host_label)

            y_row3 = header_text_y(2)
            p.drawString(x_left + 5, y_row3, "장 소")
            p.drawString(x_left + 5 + 80, y_row3, place)
            p.drawString(x_right_block + 5, y_row3, f"참석인원  {attendees_count}")

            y_row4 = header_text_y(3)
            p.drawString(x_left + 5, y_row4, "주요안건")
            if main_agenda:
                short_agenda = (main_agenda[:50] + "…") if len(main_agenda) > 50 else main_agenda
                p.drawString(x_left + 5 + 80, y_row4, short_agenda)

            current_top = header_top - 4 * header_row_h
        else:
            current_top = y

        if has_contents:
            contents_title_top = current_top
            contents_title_bottom = contents_title_top - label_row_h
            p.line(x_left, contents_title_bottom, x_right, contents_title_bottom)
            p.drawCentredString(center_x, contents_title_bottom + 8, "회의 내용")

            contents_box_top = contents_title_bottom
            contents_box_bottom = contents_box_top - contents_h
            p.line(x_left, contents_box_bottom, x_right, contents_box_bottom)

            current_top = contents_box_bottom
        else:
            contents_box_top = contents_box_bottom = None

        if has_results:
            results_title_top = current_top
            results_title_bottom = results_title_top - label_row_h
            p.line(x_left, results_title_bottom, x_right, results_title_bottom)
            p.drawCentredString(center_x, results_title_bottom + 8, "회의 결과")

            results_box_top = results_title_bottom
            results_box_bottom = results_box_top - results_h
            p.line(x_left, results_box_bottom, x_right, results_box_bottom)

            current_top = results_box_bottom
        else:
            results_box_top = results_box_bottom = None

        if has_todos:
            todos_title_top = current_top
            todos_title_bottom = todos_title_top - label_row_h
            p.line(x_left, todos_title_bottom, x_right, todos_title_bottom)
            p.drawCentredString(center_x, todos_title_bottom + 8, "해야 할 일")

            todos_box_top = todos_title_bottom
            todos_box_bottom = todos_box_top - todos_h
            p.line(x_left, todos_box_bottom, x_right, todos_box_bottom)

            current_top = todos_box_bottom
        else:
            todos_box_top = todos_box_bottom = None

        p.setFont(font, 10)

        def draw_multiline_in_box(text, x_left_box, x_right_box, top_y, bottom_y):
            max_chars = 80
            y_pos = top_y - 14
            for raw_line in (text or "").splitlines():
                line = raw_line
                while len(line) > max_chars:
                    chunk = line[:max_chars]
                    line = line[max_chars:]
                    if y_pos < bottom_y + line_height:
                        return
                    p.drawString(x_left_box + 5, y_pos, chunk)
                    y_pos -= line_height
                if line:
                    if y_pos < bottom_y + line_height:
                        return
                    p.drawString(x_left_box + 5, y_pos, line)
                    y_pos -= line_height

        if has_contents and contents_box_top is not None:
            draw_multiline_in_box(
                contents_text,
                x_left,
                x_right,
                contents_box_top,
                contents_box_bottom,
            )

        if has_results and results_box_top is not None:
            draw_multiline_in_box(
                results_text,
                x_left,
                x_right,
                results_box_top,
                results_box_bottom,
            )

        if has_todos and todos_box_top is not None:
            draw_multiline_in_box(
                todos_text,
                x_left,
                x_right,
                todos_box_top,
                todos_box_bottom,
            )

        p.setFont(font, 11)

        if has_attendees_section or not sections:
            attend_label_top = current_top
            attend_label_bottom = attend_label_top - label_row_h
            attend_block_height = label_row_h + attend_header_h + attend_rows_h

            if attend_label_top - attend_block_height < bottom_margin:
                outer_bottom = current_top
                finalize_current_page(outer_bottom)
                p.showPage()
                start_new_page()
                attend_label_top = current_top
                attend_label_bottom = attend_label_top - label_row_h

            p.line(x_left, attend_label_bottom, x_right, attend_label_bottom)
            p.drawCentredString(center_x, attend_label_bottom + 8, "참석자")

            attend_header_top = attend_label_bottom
            attend_header_bottom = attend_header_top - attend_header_h
            p.line(x_left, attend_header_bottom, x_right, attend_header_bottom)

            col_w = (x_right - x_left) / 6.0
            x_cols = [x_left + col_w * i for i in range(7)]

            attend_table_bottom = attend_header_bottom - attend_rows_h
            for xv in x_cols:
                p.line(xv, attend_header_top, xv, attend_table_bottom)

            header_y = attend_header_bottom + 5
            p.drawString(x_cols[0] + 5, header_y, "소 속")
            p.drawString(x_cols[1] + 5, header_y, "성 명")
            p.drawString(x_cols[2] + 5, header_y, "서 명")
            p.drawString(x_cols[3] + 5, header_y, "소 속")
            p.drawString(x_cols[4] + 5, header_y, "성 명")
            p.drawString(x_cols[5] + 5, header_y, "서 명")

            row_top = attend_header_bottom
            for _ in range(rows_per_side):
                row_top -= attend_row_h
                p.line(x_left, row_top, x_right, row_top)

            outer_bottom = attend_table_bottom
        else:
            outer_bottom = current_top

        finalize_current_page(outer_bottom)

        if has_attendees_section or not sections:
            attendees_list = list(attendees_qs)
            rows_per_side = max(4, math.ceil(len(attendees_list) / 2))
            col_w = (x_right - x_left) / 6.0
            x_cols = [x_left + col_w * i for i in range(7)]
            for row_idx in range(rows_per_side):
                row_text_y = attend_header_bottom - attend_row_h * row_idx - attend_row_h + 5

                left_idx = row_idx
                if left_idx < len(attendees_list):
                    att = attendees_list[left_idx]
                    dept = att.user.dept.dept_name if att.user.dept else ""
                    name = att.user.name
                    p.drawString(x_cols[0] + 5, row_text_y, dept)
                    p.drawString(x_cols[1] + 5, row_text_y, name)
                p.drawString(x_cols[2] + 5, row_text_y, "(인)")

                right_idx = row_idx + rows_per_side
                if right_idx < len(attendees_list):
                    att = attendees_list[right_idx]
                    dept = att.user.dept.dept_name if att.user.dept else ""
                    name = att.user.name
                    p.drawString(x_cols[3] + 5, row_text_y, dept)
                    p.drawString(x_cols[4] + 5, row_text_y, name)
                p.drawString(x_cols[5] + 5, row_text_y, "(인)")

        p.showPage()
        p.save()

        pdf_value = buffer.getvalue()
        buffer.close()

        filename = f"meeting_{meeting_id}_minutes.pdf"
        response = HttpResponse(pdf_value, content_type="application/pdf")
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response

    if fmt == "docx":
        doc = Document()
        doc.add_heading(meeting.title or "회의록", level=1)

        info_rows = [
            ("안건", agenda or "-"),
            ("일시", meeting_dt_str or "-"),
            ("장소", place or "-"),
            ("주최자", host_name or "-"),
            ("주요안건", main_agenda or "-"),
            ("참석자 수", str(attendees_count)),
        ]
        info_table = doc.add_table(rows=len(info_rows), cols=2)
        info_table.style = "Table Grid"
        for idx, (label, value) in enumerate(info_rows):
            cells = info_table.rows[idx].cells
            cells[0].text = label
            cells[1].text = value

        def add_doc_section(title: str, text: str):
            cleaned = (text or "").strip()
            if not cleaned:
                return
            doc.add_paragraph("")
            doc.add_heading(title, level=2)
            for ln in cleaned.splitlines():
                doc.add_paragraph(ln)

        add_doc_section("회의 내용", contents_text)
        add_doc_section("회의 결과", results_text)
        add_doc_section("해야 할 일", todos_text)

        doc.add_paragraph("")
        doc.add_heading("참석자", level=2)
        attendees_list = list(attendees_qs)
        if attendees_list:
            rows = max(1, math.ceil(len(attendees_list) / 2))
            table = doc.add_table(rows=rows + 1, cols=4)
            table.style = "Table Grid"
            header_cells = table.rows[0].cells
            header_cells[0].text = "소 속"
            header_cells[1].text = "성 명"
            header_cells[2].text = "소 속"
            header_cells[3].text = "성 명"

            for row_idx in range(rows):
                row_cells = table.rows[row_idx + 1].cells
                left_idx = row_idx
                if left_idx < len(attendees_list):
                    att = attendees_list[left_idx]
                    dept = att.user.dept.dept_name if att.user.dept else ""
                    name = att.user.name
                    row_cells[0].text = dept
                    row_cells[1].text = name
                right_idx = row_idx + rows
                if right_idx < len(attendees_list):
                    att = attendees_list[right_idx]
                    dept = att.user.dept.dept_name if att.user.dept else ""
                    name = att.user.name
                    row_cells[2].text = dept
                    row_cells[3].text = name
        else:
            doc.add_paragraph("참석자 정보가 없습니다.")

        if not meeting.meeting_notes and text_body:
            doc.add_paragraph("")
            doc.add_heading("회의 내용", level=2)
            for line in text_body.splitlines():
                doc.add_paragraph(line)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)

        filename = f"meeting_{meeting_id}_minutes.docx"
        response = HttpResponse(
            buf.getvalue(),
            content_type=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response

    return HttpResponse("invalid format", status=400)


