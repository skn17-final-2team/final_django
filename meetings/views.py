from django.shortcuts import redirect, render
from django.views.generic import TemplateView
from core.views import LoginRequiredSessionMixin
from users.models import Dept, User
from django.db.models import Prefetch, Q

from django.shortcuts import get_object_or_404
from django.http import JsonResponse, HttpResponse, Http404, StreamingHttpResponse
from django.urls import reverse

from .models import Meeting, Attendee, Task, S3File
from django.contrib import messages
from django.db import transaction


from meetings.utils.s3_upload import upload_raw_file_bytes, get_presigned_url
from meetings.utils.runpod import get_stt, get_sllm

from django.views.decorators.http import require_GET, require_POST
from datetime import date, datetime, timedelta
from django.utils import timezone

from io import BytesIO
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from django.utils.html import strip_tags, escape
import json
import math
import re
import ast
from typing import Dict
import requests
from urllib.parse import quote

from django.conf import settings
from pathlib import Path
import boto3
from botocore.config import Config

# 한글 폰트 등록 (맑은 고딕 사용)
KOREAN_FONT_NAME = settings.KOREAN_FONT_NAME

def _register_korean_font():
    # 이미 등록되어 있으면 바로 종료
    if KOREAN_FONT_NAME in pdfmetrics.getRegisteredFontNames():
        return

    font_path: Path = settings.KOREAN_FONT_PATH

    if not font_path.exists():
        raise FileNotFoundError(f"Korean font file not found: {font_path}")

    pdfmetrics.registerFont(TTFont(KOREAN_FONT_NAME, str(font_path)))

# 모듈 import 시 한 번 호출
_register_korean_font()

def _resolve_s3_file(meeting: Meeting):
    """
    meeting.record_url_id에는 presigned URL 또는 s3_key가 들어올 수 있다.
    두 경우 모두 S3File을 찾아 반환한다.
    """
    if getattr(meeting, "record_url", None):
        try:
            return meeting.record_url
        except Exception:
            pass

    record_url_val = getattr(meeting, "record_url_id", None)
    if not record_url_val:
        return None

    s3_obj = S3File.objects.filter(record_url=record_url_val).first()
    if s3_obj:
        return s3_obj

    return S3File.objects.filter(s3_key=record_url_val).first()

def _transcript_to_plain_text(raw_transcript: str) -> str:
    """
    DB에 저장된 transcript(plain 문자열 혹은 JSON 리스트 형태)를
    모델에 넘길 수 있는 평문으로 변환한다.
    """
    if not raw_transcript:
        return ""

    parsed = None
    try:
        parsed = json.loads(raw_transcript)
    except (ValueError, TypeError):
        try:
            parsed = ast.literal_eval(raw_transcript)
        except (ValueError, SyntaxError):
            parsed = None

    if isinstance(parsed, list):
        lines = []
        for segment in parsed:
            if isinstance(segment, dict):
                for speaker, text in segment.items():
                    lines.append(f"{speaker}: {text}")
        return "\n".join(lines)

    return str(raw_transcript)


def _normalize_tasks(full_tasks):
    """
    모델에서 내려준 태스크를 Task 모델에 저장할 수 있는 간단한 문자열 리스트로 정규화.
    """
    results = []
    if isinstance(full_tasks, str):
        for line in full_tasks.splitlines():
            line = line.strip(" -•\t")
            if line:
                results.append(line)
        return results

    if isinstance(full_tasks, list):
        for item in full_tasks:
            text = ""
            if isinstance(item, dict):
                who = item.get("who") or item.get("owner") or item.get("speaker")
                what = (
                    item.get("what")
                    or item.get("task")
                    or item.get("content")
                    or item.get("task_content")
                )
                when = item.get("when") or item.get("due_date")
                parts = [p for p in [who, what, when] if p]
                text = " - ".join(parts)
            elif isinstance(item, str):
                text = item.strip()
            if text:
                results.append(text)
    return results


def _parse_due_date(due_str: str):
    """
    문자열로 내려온 기한을 DateField에 맞게 파싱. 실패 시 None.
    """
    if not due_str or not isinstance(due_str, str):
        return None
    due_str = due_str.strip()
    # ISO 형태 우선
    try:
        return date.fromisoformat(due_str)
    except Exception:
        pass
    # yyyy.mm.dd 형태
    try:
        return datetime.strptime(due_str, "%Y.%m.%d").date()
    except Exception:
        pass
    # yyyy-mm-dd 형태
    try:
        return datetime.strptime(due_str, "%Y-%m-%d").date()
    except Exception:
        pass
    return None


def _extract_structured_tasks(full_tasks):
    """
    full_tasks(list/dict/str)에서 description/assignee/due를 뽑아낸 리스트 반환.
    """
    tasks = []
    # 문자열 JSON이면 파싱
    if isinstance(full_tasks, str):
        s = full_tasks.strip()
        # 코드펜스 ```json ... ``` 형태 제거
        if s.startswith("```") and s.endswith("```"):
            s = re.sub(r"^```[a-zA-Z0-9]*\s*", "", s)
            s = re.sub(r"\s*```$", "", s)
        try:
            full_tasks = json.loads(s)
        except Exception:
            full_tasks = full_tasks
    # {"tasks": [...]} 형태 처리
    if isinstance(full_tasks, dict) and "tasks" in full_tasks:
        full_tasks = full_tasks.get("tasks")
    if isinstance(full_tasks, list):
        for item in full_tasks:
            if isinstance(item, dict):
                desc = (
                    item.get("description")
                    or item.get("what")
                    or item.get("task")
                    or item.get("content")
                    or item.get("task_content")
                )
                assignee_name = (
                    item.get("assignee")
                    or item.get("who")
                    or item.get("owner")
                    or item.get("speaker")
                )
                # due/due_text 우선, due_date는 별도 보존
                due_raw = (
                    item.get("due")
                    or item.get("when")
                    or item.get("due_text")
                )
                if isinstance(due_raw, str) and due_raw.strip().lower() in ("*", "null", "none", ""):
                    due_raw = ""
                due_date_val = item.get("due_date")
                parsed_due_date = None
                if due_raw:
                    parsed_due_date = _parse_due_date(due_raw)
                elif due_date_val and str(due_date_val).lower() not in ("null", "none"):
                    parsed_due_date = _parse_due_date(str(due_date_val))
                tasks.append(
                    {
                        "description": desc if desc is not None else "",
                        "assignee_name": assignee_name if assignee_name else "",
                        "due_date": parsed_due_date,
                        "due_raw": due_raw or "",
                    }
                )
            elif isinstance(item, str):
                tasks.append(
                    {
                        "description": item,
                        "assignee_name": "",
                        "due_date": None,
                        "due_raw": "",
                    }
                )
    elif isinstance(full_tasks, str) and full_tasks.strip():
        tasks.append(
            {
                "description": full_tasks.strip(),
                "assignee_name": "",
                "due_date": None,
                "due_raw": "",
            }
        )
    return tasks


def _normalize_summary_text(full_summary):
    """
    full_summary가 문자열(코드펜스 포함 가능) 또는 dict/list로 올 때
    구조를 최대한 유지한 JSON 문자열로 반환.
    """
    if full_summary is None:
        return ""
    # 문자열이면 코드펜스 제거 후 json 파싱 시도
    if isinstance(full_summary, str):
        s = full_summary.strip()
        if s.startswith("```") and s.endswith("```"):
            s = re.sub(r"^```[a-zA-Z0-9]*\s*", "", s)
            s = re.sub(r"\s*```$", "", s)
        try:
            parsed = json.loads(s)
            if isinstance(parsed, (dict, list)):
                return json.dumps(parsed, ensure_ascii=False, indent=2)
            return s
        except Exception:
            return s
    if isinstance(full_summary, (dict, list)):
        try:
            return json.dumps(full_summary, ensure_ascii=False, indent=2)
        except Exception:
            return str(full_summary)
    return str(full_summary)


def _stringify_agenda_summary(summary_val):
    """
    Agenda summary가 dict로 내려올 때 보기 좋은 문자열로 변환한다.
    - agenda_summary 필드를 우선 사용
    - 그 외 주요 필드(who/what/when/where/why/how/how_much/how_many)를 라벨과 함께 표시
    """
    if summary_val is None:
        return ""
    if isinstance(summary_val, str):
        return summary_val.strip()
    if isinstance(summary_val, dict):
        agenda_summary = summary_val.get("agenda_summary")
        if agenda_summary:
            return str(agenda_summary).strip()
        parts = []
        label_map = {
            "who": "who",
            "what": "what",
            "when": "when",
            "where": "where",
            "why": "why",
            "how": "how",
            "how_much": "how_much",
            "how_many": "how_many",
        }
        for key in ["who", "what", "when", "where", "why", "how", "how_much", "how_many"]:
            val = summary_val.get(key)
            if val:
                parts.append(f"{label_map[key]}: {val}")
        if parts:
            return " | ".join(parts)
        try:
            return json.dumps(summary_val, ensure_ascii=False)
        except Exception:
            return str(summary_val)
    try:
        return json.dumps(summary_val, ensure_ascii=False)
    except Exception:
        return str(summary_val)


def _parse_summary_agendas(summary_text):
    """
    저장된 summary_text(JSON 문자열 예상)에서 agendas 배열을 파싱해 리스트 반환.
    """
    if not summary_text:
        return []
    parsed = None
    try:
        text = summary_text
        if isinstance(summary_text, str):
            text = summary_text.strip()
            # 코드펜스 제거
            if text.startswith("```") and text.endswith("```"):
                text = re.sub(r"^```[a-zA-Z0-9]*\s*", "", text)
                text = re.sub(r"\s*```$", "", text)
        parsed = json.loads(text)
    except Exception:
        # dict 자체로 저장되었거나, 문자열 파싱이 실패한 경우 원본을 그대로 사용
        parsed = summary_text if isinstance(summary_text, dict) else None
    if not isinstance(parsed, dict):
        return []
    agendas = parsed.get("agendas")
    if not isinstance(agendas, list):
        return []
    result = []
    for item in agendas:
        if not isinstance(item, dict):
            continue
        result.append(
            {
                "agenda": (item.get("agenda") or "").strip(),
                "agenda_description": (item.get("agenda_description") or "").strip(),
                "summary": _stringify_agenda_summary(item.get("summary")),
            }
        )
    return result


def _format_agendas_for_minutes(agendas):
    """
    회의록 기본값으로 넣기 위한 간단한 텍스트 형태로 변환.
    """
    if not agendas:
        return ""
    lines = []
    for idx, ag in enumerate(agendas, 1):
        title = ag.get("agenda") or ""
        desc = ag.get("agenda_description") or ""
        summ = _stringify_agenda_summary(ag.get("summary"))
        parts = [p for p in [title, desc, summ] if p]
        if parts:
            lines.append(f"{idx}. " + " | ".join(parts))
    return "\n".join(lines)

def _get_privacy_from_domain(domain_value):
    """
    기존 BooleanField(domain) 호환 + legacy 문자열 처리.
    """
    if domain_value in [True, False]:
        return "private" if domain_value else "public"
    try:
        if isinstance(domain_value, str):
            parsed = json.loads(domain_value)
            if isinstance(parsed, dict):
                val = parsed.get("privacy")
                if isinstance(val, str) and val.lower() == "private":
                    return "private"
            if domain_value.strip().lower() == "private":
                return "private"
    except Exception:
        pass
    return "public"

def _get_privacy(meeting: Meeting):
    if getattr(meeting, "private_yn", False):
        return "private"
    return _get_privacy_from_domain(getattr(meeting, "domain", None))


def _has_meeting_view_permission(
    meeting: Meeting,
    attendees_list,
    session_user_id,
    login_user_dept_id,
):
    """
    상세/다운로드 공통 접근 권한 체크.
    """
    is_host = session_user_id and str(meeting.host_id) == str(session_user_id)
    is_attendee = (
        session_user_id
        and any(str(a.user_id) == str(session_user_id) for a in attendees_list)
    )
    privacy = _get_privacy(meeting)
    if privacy == "private":
        return bool(is_host or is_attendee)

    same_dept = False
    if login_user_dept_id:
        same_dept = any(
            getattr(a.user, "dept_id", None) == login_user_dept_id
            for a in attendees_list
        )
    return bool(is_host or is_attendee or same_dept)


def _task_to_display(task):
    """
    Task 객체에서 화면에 표시할 who/what/when 정보를 추출.
    task_content 이 JSON(dict) 형태이면 description/due_text/assignee 사용, 아니면 문자열 그대로.
    """
    what = task.task_content or ""
    when_text = ""
    who_text_from_json = ""
    assignee_id_from_json = None
    try:
        parsed = json.loads(task.task_content)
        if isinstance(parsed, dict):
            if "tasks" in parsed and isinstance(parsed["tasks"], list) and parsed["tasks"]:
                first = parsed["tasks"][0]
                what = first.get("description") or what
                when_text = first.get("due") or first.get("due_text") or ""
                who_text_from_json = first.get("assignee") or ""
                assignee_id_from_json = first.get("assignee_id")
            else:
                what = parsed.get("description") or what
                when_text = parsed.get("due") or parsed.get("due_text") or ""
                who_text_from_json = parsed.get("assignee") or ""
                assignee_id_from_json = parsed.get("assignee_id")
    except Exception:
        pass

    who_text = "직접입력"
    if task.assignee:
        who_text = task.assignee.name
        if getattr(task.assignee, "dept", None):
            who_text += f" ({task.assignee.dept.dept_name})"
    elif who_text_from_json:
        who_text = who_text_from_json

    return {
        "id": getattr(task, "task_id", None),
        "who": who_text,
        "what": what,
        "when": when_text or "직접입력",
        "assignee_id": task.assignee_id or assignee_id_from_json,
    }


def _plain_text_to_structured(text: str):
    """
    '발화자: 내용' 형태의 평문을 구조화 목록으로 변환한다.
    """
    segments = []
    speakers = set()
    if not text:
        return segments, speakers
    for line in text.splitlines():
        line = line.strip()
        if not line or ":" not in line:
            continue
        speaker, value = line.split(":", 1)
        speaker = speaker.strip()
        value = value.strip()
        if speaker and value:
            segments.append({speaker: value})
            speakers.add(speaker)
    return segments, speakers


def _render_transcript_html(raw_transcript: str) -> str:
    """
    상세 화면에서 쓸 전문 HTML을 생성한다.
    - 리스트(JSON/리터럴)인 경우: 각 발화자:내용을 굵게 표시
    - 평문에서도 '화자: 내용' 패턴이면 굵게 표시
    """
    if not raw_transcript:
        return ""

    html_lines = []
    parsed = None
    try:
        parsed = json.loads(raw_transcript)
    except (ValueError, TypeError):
        try:
            parsed = ast.literal_eval(raw_transcript)
        except (ValueError, SyntaxError):
            parsed = None

    def add_line(speaker, text):
        html_lines.append(f"<div><strong>{escape(str(speaker))}</strong>: {escape(str(text))}</div>")

    if isinstance(parsed, list):
        for segment in parsed:
            if not isinstance(segment, dict):
                continue
            for speaker, text in segment.items():
                add_line(speaker, text)
        return "".join(html_lines)

    # 평문 처리
    for line in raw_transcript.splitlines():
        if ":" in line:
            sp, txt = line.split(":", 1)
            add_line(sp.strip(), txt.strip())
        else:
            html_lines.append(f"<div>{escape(line)}</div>")
    return "".join(html_lines)


# 회의 목록에서 쓸 데이터 생성하는 함수
def build_meeting_list_context(meeting_qs, login_user_id=None):
    meetings_data = []

    # 로그인 사용자 객체(필요하면)
    login_user = None
    login_user_dept_id = None
    if login_user_id:
        login_user = User.objects.select_related("dept").get(user_id=login_user_id)
        login_user_dept_id = getattr(login_user, "dept_id", None)

    for m in meeting_qs:
        attendees = list(m.attendees.all())
        attendee_count = len(attendees)
        attendee_names = ", ".join(a.user.name for a in attendees)

        # 참여 여부 (All 페이지에서만 실제로 사용, Mine/Dept 에선 옵션)
        is_joined = False
        is_host = str(m.host_id) == str(login_user_id) if login_user_id else False
        is_attendee = False
        if login_user_id:
            if is_host:
                is_joined = True
            else:
                is_attendee = any(str(a.user_id) == str(login_user_id) for a in attendees)
                is_joined = is_attendee

        # 공개 여부/접근 권한
        privacy = _get_privacy(m)
        if privacy == "private":
            allowed = is_host or is_attendee
        else:
            same_dept = False
            if login_user_dept_id:
                same_dept = any(
                    getattr(a.user, "dept_id", None) == login_user_dept_id
                    for a in attendees
                )
            allowed = is_host or is_attendee or same_dept
        if not allowed:
            continue

        meetings_data.append(
            {
                "meeting_id": m.meeting_id,
                "title": m.title,
                "meet_date_time": m.meet_date_time,
                "place": m.place,
                "host_name": m.host.name,
                "attendee_count": attendee_count,
                "attendee_names": attendee_names,
                "is_joined": is_joined,
                "is_private": privacy == "private",
            }
        )

    return meetings_data, login_user

class MeetingListAllView(LoginRequiredSessionMixin, TemplateView):
    template_name = "meetings/meeting_list.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        login_user_id = self.request.session.get("login_user_id")

        meeting_qs = (
            Meeting.objects
            .select_related("host")
            .prefetch_related(
                Prefetch(
                    "attendees",
                    queryset=Attendee.objects.select_related("user", "user__dept"),
                )
            )
            .order_by("-meet_date_time")
        )

        meetings_data, login_user = build_meeting_list_context(
            meeting_qs, login_user_id
        )

        context["login_user"] = login_user
        context["meetings"] = meetings_data
        context["meeting_list_type"] = "all"
        return context


class MeetingListMineView(LoginRequiredSessionMixin, TemplateView):
    template_name = "meetings/meeting_list.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        context["meeting_list_type"] = "mine"
        login_user_id = self.request.session.get("login_user_id")
        if not login_user_id:
            # 이 View는 LoginRequiredMixin 이라 실제로는 여기 안 오겠지만 방어 코드
            context["login_user"] = None
            context["meetings"] = []
            return context

        meeting_qs = (
            Meeting.objects
            .select_related("host")
            .prefetch_related(
                Prefetch(
                    "attendees",
                    queryset=Attendee.objects.select_related("user", "user__dept"),
                )
            )
            .filter(
                Q(host_id=login_user_id) | Q(attendees__user_id=login_user_id)
            )
            .order_by("-meet_date_time")
            .distinct()
        )

        meetings_data, login_user = build_meeting_list_context(
            meeting_qs, login_user_id
        )

        context["login_user"] = login_user
        context["meetings"] = meetings_data
        return context

class MeetingListDeptView(LoginRequiredSessionMixin, TemplateView):
    template_name = "meetings/meeting_list.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        # 1) 로그인 사용자
        login_user_id = self.request.session.get("login_user_id")
        login_user = None
        context["meeting_list_type"] = "dept"
        if login_user_id:
            login_user = User.objects.select_related("dept").get(user_id=login_user_id)

        # 2) 열람 가능한 회의:
        #    host도 아니고, attendee_tbl에도 없는 회의만(즉, 부서원이 참여한 회의임)
        #    단, 해당 회의는 '참석자들의 부서원' 또는 '주최자의 부서원'인 경우에만 열람 가능
        dept_id = getattr(login_user, "dept_id", None)
        if dept_id:
            meeting_qs = (
                Meeting.objects
                .select_related("host")
                .prefetch_related(
                    Prefetch(
                        "attendees",
                        queryset=Attendee.objects.select_related("user", "user__dept"),
                    )
                )
                .exclude(
                    Q(host_id=login_user_id) | Q(attendees__user_id=login_user_id)
                )
                .filter(
                    Q(attendees__user__dept_id=dept_id) | Q(host__dept_id=dept_id)
                )
                .order_by("-meet_date_time")
                .distinct()
            )
        else:
            meeting_qs = Meeting.objects.none()

        # 3) 템플릿용 데이터 가공 (전체/내가참여한 과 동일 구조)
        meetings_data = []

        for m in meeting_qs:
            attendees = list(m.attendees.all())
            attendee_count = len(attendees)
            attendee_names = ", ".join(a.user.name for a in attendees)

            # 이 View는 "미참여 회의"라 is_joined는 항상 False
            is_joined = False

            meetings_data.append(
                {
                    "meeting_id": m.meeting_id,
                    "title": m.title,
                    "meet_date_time": m.meet_date_time,
                    "place": m.place,
                    "host_name": m.host.name,
                    "attendee_count": attendee_count,
                    "attendee_names": attendee_names,
                    "is_joined": is_joined,
                }
            )

        context["login_user"] = login_user
        context["meetings"] = meetings_data
        return context

class MeetingCreateView(LoginRequiredSessionMixin, TemplateView):
    template_name = "meetings/meeting_create.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        # DB에서 실제 부서 + 유저들을 가져오는 부분
        dept_qs = Dept.objects.prefetch_related(
            Prefetch("users", queryset=User.objects.order_by("name"))
        ).order_by("dept_name")

        context["departments"] = dept_qs   # 템플릿으로 넘김
        # 로그인한 사용자 정보도 템플릿으로 내려줌
        login_user_id = self.request.session.get("login_user_id")
        if login_user_id:
            context["login_user"] = User.objects.get(user_id=login_user_id)
        return context

    def post(self, request, *args, **kwargs):

        attendee_ids = request.POST.getlist("attendees")
        # 도메인은 사용자가 선택한 한글 라벨 그대로 저장한다.
        domain_value = request.POST.get("domains") or ""
        privacy_private = request.POST.get("is_private") in ["on", "1", "true", True]
        if not attendee_ids:
            messages.error(request, "참석자를 최소 1명 이상 선택해 주세요.")
            context = self.get_context_data()
            context["form_title"] = request.POST.get("title", "")
            context["form_meet_date_time"] = request.POST.get("meet_date_time", "")
            context["form_place"] = request.POST.get("place", "")
            return render(request, self.template_name, context)
        
        # 1. 폼 데이터
        title = request.POST.get("title")
        meet_date_time = request.POST.get("meet_date_time")
        place = request.POST.get("place")

        # 2. 세션에서 로그인한 사용자 id 가져오기
        login_user_id = request.session.get("login_user_id")
        if not login_user_id:
            # 혹시 세션 끊긴 경우 대비
            return redirect("login")

        # 3. host로 쓸 User 객체 조회
        host_user = User.objects.get(user_id=login_user_id)
        
        attendee_ids = [uid for uid in attendee_ids if str(uid) != str(login_user_id)]


        # 4. meeting_tbl에 새 레코드 생성
        with transaction.atomic():
            # meeting_tbl insert
            meeting = Meeting.objects.create(
                title=title,
                meet_date_time=meet_date_time,
                place=place,
                host=host_user,   # FK: User 인스턴스
                transcript="",    # NOT NULL 필드라면 임시값
                domain=domain_value,
                private_yn=privacy_private,
            )

        Attendee.objects.create(meeting=meeting, user=host_user)

        users = User.objects.filter(user_id__in=attendee_ids)
        attendee_objs = [
                Attendee(meeting=meeting, user=u) for u in users
        ]
        Attendee.objects.bulk_create(attendee_objs)

        # 5. 생성된 meeting_id를 가지고 녹음 화면으로 이동
        return redirect("meetings:meeting_record", meeting_id=meeting.meeting_id)
    
class MeetingRecordView(LoginRequiredSessionMixin, TemplateView):
    template_name = "meetings/meeting_record.html"

    def get(self, request, *args, **kwargs):
        # 권한 검사: 상세 보기와 동일한 정책을 적용하여 불허 시 리다이렉트
        meeting_id = kwargs.get("meeting_id") or self.kwargs.get("meeting_id")
        meeting = (
            Meeting.objects
            .select_related("host")
            .prefetch_related("attendees__user__dept")
            .filter(pk=meeting_id)
            .first()
        )
        if not meeting:
            raise Http404()

        session_user_id = request.session.get("login_user_id")
        login_user_obj = None
        login_user_dept_id = None
        if session_user_id:
            try:
                login_user_obj = User.objects.select_related("dept").get(user_id=session_user_id)
                login_user_dept_id = getattr(login_user_obj, "dept_id", None)
            except User.DoesNotExist:
                login_user_obj = None

        attendees_list = list(meeting.attendees.all())
        is_host = session_user_id and str(meeting.host_id) == str(session_user_id)
        is_attendee = any(str(a.user_id) == str(session_user_id) for a in attendees_list) if session_user_id else False
        privacy = _get_privacy(meeting)
        if privacy == "private":
            allowed = is_host or is_attendee
        else:
            same_dept = False
            if login_user_dept_id:
                same_dept = any(
                    getattr(a.user, "dept_id", None) == login_user_dept_id
                    for a in attendees_list
                )
            allowed = is_host or is_attendee or same_dept
        if not allowed:
            referer = request.META.get("HTTP_REFERER")
            if referer:
                return redirect(referer)
            try:
                fallback = reverse("meetings:meeting_list_dept")
            except Exception:
                fallback = "/"
            return redirect(fallback)

        return super().get(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        meeting_id = self.kwargs.get("meeting_id")

        meeting = Meeting.objects.get(pk=meeting_id)

        context["attendees"] = (
            meeting.attendees
                   .select_related("user")  # Attendee.user
                   .all()
        )

        # 녹음 허용 여부: 회의 시작 시간이 현재 시점보다 과거여도
        # 허용 여유(grace) 내라면 녹음 UI를 노출합니다.
        grace = timedelta(minutes=10)
        now = timezone.now()
        allow_recording = False
        try:
            if meeting.meet_date_time:
                allow_recording = meeting.meet_date_time >= (now - grace)
        except Exception:
            allow_recording = False

        context["allow_recording"] = allow_recording
        context["meeting"] = meeting
        context["meeting_id"] = meeting_id
        return context

class MeetingTranscriptView(LoginRequiredSessionMixin, TemplateView):
    template_name = "meetings/meeting_transcript.html"

    def get(self, request, *args, **kwargs):
        # 권한 검사: 상세 보기와 동일한 정책을 적용하여 불허 시 리다이렉트
        meeting_id = kwargs.get("meeting_id") or self.kwargs.get("meeting_id")
        meeting = (
            Meeting.objects
            .select_related("host")
            .prefetch_related("attendees__user__dept")
            .filter(pk=meeting_id)
            .first()
        )
        if not meeting:
            raise Http404()

        session_user_id = request.session.get("login_user_id")
        login_user_obj = None
        login_user_dept_id = None
        if session_user_id:
            try:
                login_user_obj = User.objects.select_related("dept").get(user_id=session_user_id)
                login_user_dept_id = getattr(login_user_obj, "dept_id", None)
            except User.DoesNotExist:
                login_user_obj = None

        attendees_list = list(meeting.attendees.all())
        is_host = session_user_id and str(meeting.host_id) == str(session_user_id)
        is_attendee = any(str(a.user_id) == str(session_user_id) for a in attendees_list) if session_user_id else False
        privacy = _get_privacy(meeting)
        if privacy == "private":
            allowed = is_host or is_attendee
        else:
            same_dept = False
            if login_user_dept_id:
                same_dept = any(
                    getattr(a.user, "dept_id", None) == login_user_dept_id
                    for a in attendees_list
                )
            allowed = is_host or is_attendee or same_dept
        if not allowed:
            referer = request.META.get("HTTP_REFERER")
            if referer:
                return redirect(referer)
            try:
                fallback = reverse("meetings:meeting_list_dept")
            except Exception:
                fallback = "/"
            return redirect(fallback)

        return super().get(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        meeting_id = self.kwargs.get("meeting_id")
        context["meeting_id"] = meeting_id
        context["meeting"] = get_object_or_404(Meeting, pk=meeting_id)
        return context

@require_GET
def meeting_transcript_api(request, meeting_id):
    meeting = get_object_or_404(Meeting, pk=meeting_id)
    attendees_qs = meeting.attendees.select_related("user", "user__dept").all()

    raw_transcript = meeting.transcript or ""
    structured_transcript = []
    speakers = []
    transcript_plain = raw_transcript.replace("\r\n", "\n")

    if raw_transcript:
        parsed = None
        try:
            parsed = json.loads(raw_transcript)
        except (ValueError, TypeError):
            try:
                parsed = ast.literal_eval(raw_transcript)
            except (ValueError, SyntaxError):
                parsed = None

        if isinstance(parsed, list):
            speaker_keys = set()
            plain_lines = []
            normalized_segments = []

            for segment in parsed:
                if not isinstance(segment, dict):
                    continue
                normalized_segment = {}
                for key, value in segment.items():
                    key_str = str(key)
                    value_str = str(value)
                    normalized_segment[key_str] = value_str
                    speaker_keys.add(key_str)
                    plain_lines.append(f"{key_str}: {value_str}")
                if normalized_segment:
                    normalized_segments.append(normalized_segment)

            if normalized_segments:
                structured_transcript = normalized_segments
                speakers = sorted(speaker_keys)
                transcript_plain = "\n".join(plain_lines)
        # else: keep defaults (plain text only)

    attendees_payload = [
        {
            "user_id": attendee.user_id,
            "name": attendee.user.name,
            "dept_name": attendee.user.dept.dept_name if attendee.user.dept else "",
        }
        for attendee in attendees_qs
    ]

    return JsonResponse(
        {
            "meeting_title": meeting.title,
            "transcript": raw_transcript,
            "transcript_plain": transcript_plain,
            "transcript_structured": structured_transcript,
            "record_url": str(meeting.record_url_id) if meeting.record_url_id else "",
            "attendees": attendees_payload,
            "speakers": speakers,
        }
    )


@require_POST
def meeting_transcript_save(request, meeting_id):
    meeting = get_object_or_404(Meeting, pk=meeting_id)

    session_user_id = request.session.get("login_user_id")
    request_user_id = None
    if hasattr(request, "user") and request.user.is_authenticated:
        request_user_id = getattr(request.user, "user_id", None) or request.user.id

    host_user_id = str(meeting.host_id) if meeting.host_id else None
    has_permission = False
    if host_user_id:
        if session_user_id and str(session_user_id) == host_user_id:
            has_permission = True
        elif request_user_id and str(request_user_id) == host_user_id:
            has_permission = True

    if not has_permission:
        return JsonResponse(
            {"ok": False, "error": "전문을 저장할 권한이 없습니다."},
            status=403,
        )

    try:
        payload = json.loads(request.body.decode("utf-8"))
    except json.JSONDecodeError:
        return JsonResponse({"ok": False, "error": "Invalid JSON"}, status=400)

    transcript_text = (payload.get("transcript_text") or "").strip()
    transcript_structured = payload.get("transcript_structured")

    if transcript_structured:
        if not isinstance(transcript_structured, list):
            return JsonResponse(
                {"ok": False, "error": "잘못된 전문 형식입니다."},
                status=400,
            )
        try:
            meeting.transcript = json.dumps(transcript_structured, ensure_ascii=False)
        except (TypeError, ValueError):
            return JsonResponse(
                {"ok": False, "error": "전문 데이터를 처리할 수 없습니다."},
                status=400,
            )
    elif transcript_text:
        meeting.transcript = transcript_text
    else:
        return JsonResponse(
            {"ok": False, "error": "저장할 전문 내용이 없습니다."},
            status=400,
        )

    meeting.save(update_fields=["transcript"])

    redirect_url = reverse("meetings:rendering_sllm", args=[meeting_id])
    return JsonResponse({"ok": True, "redirect_url": redirect_url})



class MeetingDetailView(LoginRequiredSessionMixin, TemplateView):
    template_name = "meetings/meeting_detail.html"

    def get(self, request, *args, **kwargs):
        # 권한 검사 로직을 get()에서 처리하여 리다이렉트 응답을 반환하도록 함
        meeting_id = kwargs.get("meeting_id") or self.kwargs.get("meeting_id")
        meeting = (
            Meeting.objects
            .select_related("host")
            .prefetch_related("attendees__user__dept")
            .filter(pk=meeting_id)
            .first()
        )
        if not meeting:
            raise Http404()

        session_user_id = request.session.get("login_user_id")
        login_user_obj = None
        login_user_dept_id = None
        if session_user_id:
            try:
                login_user_obj = User.objects.select_related("dept").get(user_id=session_user_id)
                login_user_dept_id = getattr(login_user_obj, "dept_id", None)
            except User.DoesNotExist:
                login_user_obj = None

        attendees_list = list(meeting.attendees.all())
        is_host = session_user_id and str(meeting.host_id) == str(session_user_id)
        is_attendee = any(str(a.user_id) == str(session_user_id) for a in attendees_list) if session_user_id else False
        privacy = _get_privacy(meeting)
        if privacy == "private":
            allowed = is_host or is_attendee
        else:
            same_dept = False
            if login_user_dept_id:
                same_dept = any(
                    getattr(a.user, "dept_id", None) == login_user_dept_id
                    for a in attendees_list
                )
            allowed = is_host or is_attendee or same_dept
        if not allowed:
            referer = request.META.get("HTTP_REFERER")
            if referer:
                return redirect(referer)
            try:
                fallback = reverse("meetings:meeting_list_dept")
            except Exception:
                fallback = "/"
            return redirect(fallback)

        return super().get(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        meeting_id = self.kwargs.get("meeting_id")
        meeting = (
            Meeting.objects
            .select_related("host")
            .prefetch_related("attendees__user__dept")
            .filter(pk=meeting_id)
            .first()
        )
        if not meeting:
            raise Http404()

        session_user_id = self.request.session.get("login_user_id")
        login_user_obj = None
        login_user_dept_id = None
        if session_user_id:
            try:
                login_user_obj = User.objects.select_related("dept").get(user_id=session_user_id)
                login_user_dept_id = getattr(login_user_obj, "dept_id", None)
            except User.DoesNotExist:
                login_user_obj = None

        attendees_list = list(meeting.attendees.all())
        is_host = session_user_id and str(meeting.host_id) == str(session_user_id)
        is_attendee = any(str(a.user_id) == str(session_user_id) for a in attendees_list) if session_user_id else False
        privacy = _get_privacy(meeting)
        if privacy == "private":
            allowed = is_host or is_attendee
        else:
            same_dept = False
            if login_user_dept_id:
                same_dept = any(
                    getattr(a.user, "dept_id", None) == login_user_dept_id
                    for a in attendees_list
                )
            allowed = is_host or is_attendee or same_dept
        if not allowed:
            # 이전 페이지가 있으면 그쪽으로 리다이렉트, 없으면 부서 회의 목록으로 이동
            referer = self.request.META.get("HTTP_REFERER")
            if referer:
                return redirect(referer)
            try:
                fallback = reverse("meetings:meeting_list_dept")
            except Exception:
                fallback = "/"
            return redirect(fallback)

        # 템플릿에서 사용할 데이터 주입
        context["meeting"] = meeting
        context["login_user"] = login_user_obj
        context["login_user_id"] = session_user_id
        context["attendees"] = (
            meeting.attendees
                   .select_related("user", "user__dept")
                   .all()
        )
        context["tasks"] = (
            meeting.tasks
                   .select_related("assignee", "assignee__dept")
                   .all()
        )
        context["tasks_display"] = [_task_to_display(t) for t in context["tasks"]]
        # summary JSON(agendas) 파싱
        summary_agendas = _parse_summary_agendas(meeting.summary)
        context["summary_agendas"] = summary_agendas
        context["summary_agendas_text"] = _format_agendas_for_minutes(summary_agendas)
        context["summary_agenda_first"] = summary_agendas[0].get("agenda") if summary_agendas else ""
        # 해야 할 일 기본값 (회의록용)
        tasks_for_minutes = []
        for t in context["tasks_display"]:
            parts = []
            if t.get("who"):
                parts.append(f"Who: {t['who']}")
            if t.get("what"):
                parts.append(f"What: {t['what']}")
            if t.get("when"):
                parts.append(f"When: {t['when']}")
            if parts:
                tasks_for_minutes.append(" | ".join(parts))
        context["tasks_for_minutes"] = "\n".join(tasks_for_minutes)
        # 전체 사용자 목록 (who 자동완성용)
        context["all_users"] = list(
            User.objects
                .select_related("dept")
                .values("user_id", "name", "dept__dept_name")
                .order_by("name")
        )
        context["transcript_display_html"] = _render_transcript_html(meeting.transcript)
        try:
            context["all_users_json"] = json.dumps(context["all_users"], ensure_ascii=False)
        except Exception:
            context["all_users_json"] = "[]"

        # 주최자인지 여부 판단 (세션 기반)
        is_host = False

        if session_user_id and meeting.host_id:
            is_host = str(meeting.host_id) == str(session_user_id)
        elif hasattr(self.request, "user") and self.request.user.is_authenticated:
            # Django 기본 User를 통해 접근 가능한 경우까지 보조 처리
            request_user_id = getattr(self.request.user, "user_id", None) or self.request.user.id
            if request_user_id:
                is_host = str(meeting.host_id) == str(request_user_id)

        context["is_host"] = is_host          # 템플릿에서 사용
        context["can_edit_minutes"] = is_host # 이름 하나 더 두고 싶으면
        return context


@require_GET
def meeting_audio_download(request, meeting_id):
    meeting = (
        Meeting.objects.select_related("record_url", "host")
        .prefetch_related("attendees__user__dept")
        .filter(pk=meeting_id)
        .first()
    )
    if not meeting:
        raise Http404()

    session_user_id = request.session.get("login_user_id")
    login_user_dept_id = None
    if session_user_id:
        login_user_obj = (
            User.objects.select_related("dept")
            .filter(user_id=session_user_id)
            .first()
        )
        if login_user_obj:
            login_user_dept_id = getattr(login_user_obj, "dept_id", None)

    attendees_list = list(meeting.attendees.all())
    if not _has_meeting_view_permission(
        meeting,
        attendees_list,
        session_user_id,
        login_user_dept_id,
    ):
        return JsonResponse(
            {"ok": False, "error": "음성 파일을 다운로드할 권한이 없습니다."},
            status=403,
        )

    if not meeting.record_url_id:
        return JsonResponse(
            {"ok": False, "error": "등록된 음성 파일이 없습니다."},
            status=404,
        )

    s3_obj = _resolve_s3_file(meeting)
    if not s3_obj:
        return JsonResponse(
            {"ok": False, "error": "등록된 음성 파일이 없습니다."},
            status=404,
        )

    try:
        presigned_url = get_presigned_url(s3_obj.s3_key)
    except Exception:
        return JsonResponse(
            {"ok": False, "error": "음성 파일 URL을 생성하는 중 오류가 발생했습니다."},
            status=500,
        )
    try:
        s3_response = requests.get(presigned_url, stream=True, timeout=30)
    except requests.RequestException:
        return JsonResponse(
            {"ok": False, "error": "음성 파일을 가져오는 중 오류가 발생했습니다."},
            status=502,
        )

    if s3_response.status_code != 200:
        s3_response.close()
        return JsonResponse(
            {"ok": False, "error": "음성 파일을 가져오는 데 실패했습니다."},
            status=502,
        )

    def stream_file():
        try:
            for chunk in s3_response.iter_content(chunk_size=8192):
                if chunk:
                    yield chunk
        finally:
            s3_response.close()

    filename = getattr(s3_obj, "original_name", None) or f"meeting_{meeting_id}.wav"
    fallback_filename = "meeting_audio.wav"
    quoted_filename = quote(filename)
    content_type = s3_response.headers.get("Content-Type") or "audio/wav"
    content_length = s3_response.headers.get("Content-Length")

    response = StreamingHttpResponse(stream_file(), content_type=content_type)
    if content_length:
        response["Content-Length"] = content_length
    response["Content-Disposition"] = (
        f"attachment; filename*=UTF-8''{quoted_filename}; filename=\"{fallback_filename}\""
    )
    return response

def meeting_record_upload(request, meeting_id):
    # 1) 메소드 체크
    if request.method != "POST":
        return JsonResponse({"error": "Invalid method"}, status=405)

    # 2) 회의 조회
    meeting = get_object_or_404(Meeting, pk=meeting_id)

    # 3) 파일 추출
    uploaded_file = request.FILES.get("file")
    if not uploaded_file:
        return JsonResponse({"error": "file is required"}, status=400)

    # 4) 확장자 검증 (wav)
    filename = uploaded_file.name
    ext = filename.split(".")[-1].lower()
    if ext not in ["wav"]:
        return JsonResponse(
            {"error": "WAV 파일만 업로드 가능합니다."},
            status=400,
        )

    # 5) 유틸 호출
    file_bytes = uploaded_file.read()
    try:
        record_url = upload_raw_file_bytes(
            file_bytes=file_bytes,
            original_filename=filename,
            delete_after_seconds=172800, # 48시간 뒤 삭제
            # delete_after_seconds=3600, # 테스트용 1시간 뒤 삭제
        )
    except Exception as e:
        # 유틸 호출 중 에러가 나도 반드시 응답을 반환
        return JsonResponse(
            {"error": f"S3 업로드 중 오류: {str(e)}"},
            status=500,
        )

    # 6) Meeting FK 연결 (record_url 이 ForeignKey(S3File, db_column="record_url") 일 때)
    meeting.record_url_id = record_url
    meeting.save(update_fields=["record_url"])

    return JsonResponse({
        "ok": True,
        "record_url": record_url,
    })


@require_POST
def meeting_record_url_set(request, meeting_id):
    """
    업로드가 아닌 이미 존재하는 S3 키를 회의에 직접 연결할 때 사용.
    주최자만 설정 가능.
    """
    meeting = get_object_or_404(Meeting, pk=meeting_id)

    session_user_id = request.session.get("login_user_id")
    request_user_id = None
    if hasattr(request, "user") and request.user.is_authenticated:
        request_user_id = getattr(request.user, "user_id", None) or request.user.id

    host_user_id = str(meeting.host_id) if meeting.host_id else None
    has_permission = False
    if host_user_id:
        if session_user_id and str(session_user_id) == host_user_id:
            has_permission = True
        elif request_user_id and str(request_user_id) == host_user_id:
            has_permission = True

    if not has_permission:
        return JsonResponse(
            {"ok": False, "error": "녹음 파일을 설정할 권한이 없습니다."},
            status=403,
        )

    try:
        payload = json.loads(request.body.decode("utf-8"))
    except json.JSONDecodeError:
        return JsonResponse({"ok": False, "error": "Invalid JSON"}, status=400)

    s3_key = (payload.get("s3_key") or "").strip()
    if not s3_key:
        return JsonResponse({"ok": False, "error": "s3_key is required"}, status=400)

    original_name = (payload.get("original_name") or "").strip()
    if not original_name:
        original_name = s3_key.split("/")[-1] or "audio.wav"

    delete_after_seconds = payload.get("delete_after_seconds")
    if not isinstance(delete_after_seconds, int) or delete_after_seconds <= 0:
        delete_after_seconds = 60 * 60 * 24 * 30  # 기본 30일
    delete_at = timezone.now() + timedelta(seconds=delete_after_seconds)

    # S3File FK 보존을 위해 레코드가 없으면 생성/있으면 갱신
    s3_obj, created = S3File.objects.get_or_create(
        s3_key=s3_key,
        defaults={
            "original_name": original_name,
            "delete_at": delete_at,
        },
    )
    if not created:
        update_fields = []
        if original_name and s3_obj.original_name != original_name:
            s3_obj.original_name = original_name
            update_fields.append("original_name")
        # 요청마다 삭제 시점을 연장할 수 있도록 덮어쓴다.
        s3_obj.delete_at = delete_at
        update_fields.append("delete_at")
        if update_fields:
            s3_obj.save(update_fields=update_fields)

    meeting.record_url_id = s3_key
    meeting.save(update_fields=["record_url"])

    return JsonResponse({"ok": True, "s3_key": s3_key})

def meeting_summary(request, meeting_id):
    # 별도 요약 화면 대신 상세 페이지로 이동
    return redirect("meetings:meeting_detail", meeting_id=meeting_id)

class MeetingSttRenderingView(LoginRequiredSessionMixin, TemplateView):
    """
    STT 대기 화면
    """
    template_name = "meetings/rending_stt.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        meeting_id = self.kwargs.get("meeting_id")
        meeting = get_object_or_404(Meeting, pk=meeting_id)
        context["meeting"] = meeting
        context["meeting_id"] = meeting_id
        return context


class MeetingSllmRenderingView(LoginRequiredSessionMixin, TemplateView):
    """
    SLLM 대기 화면
    """
    template_name = "meetings/rending_sllm.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        meeting_id = self.kwargs.get("meeting_id")

        meeting = get_object_or_404(Meeting, pk=meeting_id)

        context["meeting"] = meeting
        context["meeting_id"] = meeting_id
        return context

@require_GET
def meeting_transcript_prepare(request, meeting_id):
    """
    렌딩 페이지에서 호출하는 엔드포인트.
    - 아직 transcript가 없으면 STT를 돌려서 생성
    - 이미 있으면 바로 done 리턴
    """
    meeting = get_object_or_404(Meeting, pk=meeting_id)

    transcript_html = meeting.transcript or ""

    if not transcript_html:
        # STT 호출
        s3_obj = _resolve_s3_file(meeting)
        if not s3_obj:
            return JsonResponse(
                {
                    "status": "error",
                    "message": "등록된 음성 파일이 없습니다.",
                },
                status=404,
            )
        try:
            presigned_url = get_presigned_url(s3_obj.s3_key)
        except Exception as e:
            return JsonResponse(
                {
                    "status": "error",
                    "message": f"음성 URL 생성 중 오류가 발생했습니다: {e}",
                },
                status=500,
            )

        print(f"[STT] meeting_id={meeting_id} s3_key={s3_obj.s3_key} presigned_url_generated")
        try:
            res = get_stt(presigned_url)
        except requests.RequestException as e:
            return JsonResponse(
                {
                    "status": "error",
                    "message": f"STT 호출 중 통신 오류가 발생했습니다: {e}",
                },
                status=502,
            )
        req_url = getattr(getattr(res, "request", None), "url", "")
        print(f"[STT][response] status={res.status_code} req_url={req_url} body_preview={getattr(res, 'text', '')[:300]}")

        if res.status_code != 200 or not res.json().get("success"):
            body_preview = ""
            try:
                body_preview = res.text[:300]
            except Exception:
                body_preview = ""

            message = None
            try:
                res_json = res.json()
                if isinstance(res_json, dict):
                    message = res_json.get("message") or res_json.get("error")
            except Exception:
                res_json = {}

            # 실패 시 프론트에서 메시지 보여줄 수 있도록 에러 내려줌
            return JsonResponse(
                {
                    "status": "error",
                    "message": message or "전사 처리 중 오류가 발생했습니다.",
                    "detail": body_preview,
                },
                status=500,
            )
        res = res.json()
        transcript_html = res['data']['full_text']
        meeting.transcript = transcript_html
        meeting.save(update_fields=["transcript"])

    # 여기까지 왔다면 transcript 는 채워진 상태
    return JsonResponse({"status": "done"})


@require_GET
def meeting_sllm_prepare(request, meeting_id):
    """
    전문(화자 매핑 완료본)을 SLLM에 전달해 요약/태스크를 생성한다.
    렌딩 페이지에서 호출하며, 완료 시 detail 화면으로 넘어간다.
    """
    meeting = get_object_or_404(Meeting, pk=meeting_id)

    transcript_plain = _transcript_to_plain_text(meeting.transcript)
    if not transcript_plain.strip():
        return JsonResponse(
            {"status": "error", "message": "분석할 전문이 없습니다."},
            status=400,
        )

    # SLLM 전달용 도메인 매핑: DB에는 한글, SLLM에는 영문 캐노니컬
    domain_raw = (meeting.domain or "").strip()
    domain_map = {
        "마케팅": "Marketing / Economy",
        "IT": "IT",
        "디자인": "Design",
        "회계": "Accounting",
    }
    domain_for_api = domain_map.get(domain_raw, domain_raw)
    domain_payload = [domain_for_api] if domain_for_api else []
    print(f"[SLLM] meeting_id={meeting_id} domain_raw='{domain_raw}' domain_for_api='{domain_for_api}' domain_payload={domain_payload}")

    try:
        res = get_sllm(transcript_plain, domain=domain_payload)
    except requests.RequestException as e:
        return JsonResponse(
            {"status": "error", "message": f"SLLM 호출 중 통신 오류가 발생했습니다: {e}"},
            status=500,
        )
    req_url = getattr(getattr(res, "request", None), "url", None)
    print("1번")
    print(f"[SLLM][request] meeting_id={meeting_id} status={getattr(res, 'status_code', None)} url={req_url}")
    try:
        res_json = res.json()
    except Exception:
        res_json = {}
    print("2번")
    print(res_json)

    payload = res_json.get("data") or res_json
    print("3번")
    print(payload)
    # 일부 응답은 success 필드를 포함하지 않을 수 있으므로, 실패 명시(false)인 경우만 실패로 간주
    if res.status_code != 200 or res_json.get("success") is False or not payload:
        body_preview = ""
        try:
            body_preview = res.text[:1000]
        except Exception:
            body_preview = ""

        # 응답 JSON에서 에러 메시지를 찾기 위해 여러 위치를 검사
        message = None
        if isinstance(res_json, dict):
            message = res_json.get("message") or res_json.get("error")
            detail = res_json.get("detail")
            if not message and isinstance(detail, dict):
                message = detail.get("message") or detail.get("error")

        # payload 안의 에러 필드도 확인
        if not message and isinstance(payload, dict):
            message = payload.get("message") or payload.get("error")

        if not message:
            message = body_preview or "SLLM 호출 중 오류가 발생했습니다."

        print(f"[SLLM][error] status={res.status_code} url={req_url} message={message}")
        return JsonResponse(
            {
                "status": "error",
                "message": message,
            },
            status=500,
        )

    full_summary_raw = payload.get("full_summary") or payload.get("summary") or ""
    # SLLM 응답이 agendas만 줄 때도 summary에 저장되도록 보완
    if not full_summary_raw and payload.get("agendas"):
        full_summary_raw = {"agendas": payload.get("agendas")}
    full_summary = _normalize_summary_text(full_summary_raw)
    full_tasks = payload.get("full_tasks") or payload.get("tasks") or []
    tasks_structured = _extract_structured_tasks(full_tasks)
    # 태스크 로그 찍기
    try:
        print(f"[SLLM] meeting_id={meeting_id} full_tasks={full_tasks}")
        print(f"[SLLM] meeting_id={meeting_id} tasks_structured={tasks_structured}")
    except Exception:
        pass

    with transaction.atomic():
        update_fields = []
        if full_summary is not None:
            meeting.summary = full_summary
            update_fields.append("summary")
        if update_fields:
            meeting.save(update_fields=update_fields)
        meeting.tasks.all().delete()
        if tasks_structured:
            task_objs = []
            for t in tasks_structured:
                desc = (t.get("description") or "")
                assignee_name = (t.get("assignee_name") or "").strip()
                due_date = t.get("due_date")
                due_raw = t.get("due_raw") or ""

                assignee_obj = None
                if assignee_name:
                    try:
                        assignee_obj = User.objects.filter(name=assignee_name).first()
                    except Exception:
                        assignee_obj = None

                content_payload = {"description": desc}
                if assignee_name:
                    content_payload["assignee"] = assignee_name
                if due_raw:
                    content_payload["due"] = due_raw
                if due_date:
                    try:
                        content_payload["due_date"] = due_date.isoformat()
                    except Exception:
                        content_payload["due_date"] = str(due_date)

                task_objs.append(
                    Task(
                        meeting=meeting,
                        task_content=json.dumps(content_payload, ensure_ascii=False),
                        assignee=assignee_obj,
                        due_date=due_date,
                    )
                )
            Task.objects.bulk_create(task_objs)

    return JsonResponse({"status": "done"})


def today_meetings(request):
    """
    모든 템플릿에서 'today_meetings'로
    meet_date_time 기준 '오늘 날짜'인 회의 목록에 접근할 수 있게 해주는 컨텍스트 프로세서
    """
    from django.utils import timezone

    today = date.today()

    login_user_id = request.session.get("login_user_id")
    login_user = None
    login_user_dept_id = None
    if login_user_id:
        try:
            login_user = User.objects.select_related("dept").get(user_id=login_user_id)
            login_user_dept_id = getattr(login_user, "dept_id", None)
        except User.DoesNotExist:
            login_user = None

    base_qs = (
        Meeting.objects
        .filter(meet_date_time__date=today)
        .select_related("host")
        .prefetch_related("attendees__user__dept")
        .order_by("-meet_date_time")
    )

    meetings = []
    now = timezone.now()

    for m in base_qs:
        attendees = list(m.attendees.all())
        is_host = login_user_id and str(m.host_id) == str(login_user_id)
        is_attendee = any(str(a.user_id) == str(login_user_id) for a in attendees) if login_user_id else False
        privacy = _get_privacy(m)
        if privacy == "private":
            allowed = is_host or is_attendee
        else:
            same_dept = False
            if login_user_dept_id:
                same_dept = any(
                    getattr(a.user, "dept_id", None) == login_user_dept_id
                    for a in attendees
                )
            allowed = is_host or is_attendee or same_dept
        if allowed:
            # 시간 차이 계산
            time_diff = now - m.meet_date_time
            total_minutes = int(time_diff.total_seconds() / 60)

            if total_minutes < 60:
                m.time_ago = f"{total_minutes}분 전"
            else:
                hours = total_minutes // 60
                m.time_ago = f"{hours}시간 전"

            meetings.append(m)
        if len(meetings) >= 3:
            break

    return {
        "today_meetings": meetings
    }

def minutes_save(request, meeting_id):
    """
    회의록 HTML(meeting_notes)을 저장하는 용도 : 주최자만 저장가능
    """
    meeting = get_object_or_404(Meeting, pk=meeting_id)

    session_user_id = request.session.get("login_user_id")
    request_user_id = None

    if hasattr(request, "user") and request.user.is_authenticated:
        request_user_id = getattr(request.user, "user_id", None) or request.user.id

    host_user_id = str(meeting.host_id) if meeting.host_id else None

    has_permission = False
    if host_user_id:
        if session_user_id and str(session_user_id) == host_user_id:
            has_permission = True
        elif request_user_id and str(request_user_id) == host_user_id:
            has_permission = True

    if not getattr(meeting, "host", None) or not has_permission:
        return JsonResponse(
            {"ok": False, "error": "회의록을 수정할 권한이 없습니다."},
            status=403,
        )

    try:
        data = json.loads(request.body.decode("utf-8"))
    except json.JSONDecodeError:
        return JsonResponse({"ok": False, "error": "Invalid JSON"}, status=400)

    content = (data.get("content") or "").strip()

    meeting.meeting_notes = content
    meeting.save(update_fields=["meeting_notes"])

    return JsonResponse({"ok": True})


@require_POST
def tasks_save(request, meeting_id):
    """
    태스크(Who/What/When) 수정본을 저장. 주최자만 가능.
    기존 태스크를 모두 삭제 후 전달된 목록으로 재생성한다.
    """
    meeting = get_object_or_404(Meeting, pk=meeting_id)

    session_user_id = request.session.get("login_user_id")
    request_user_id = None
    if hasattr(request, "user") and request.user.is_authenticated:
        request_user_id = getattr(request.user, "user_id", None) or request.user.id

    host_user_id = str(meeting.host_id) if meeting.host_id else None
    has_permission = False
    if host_user_id:
        if session_user_id and str(session_user_id) == host_user_id:
            has_permission = True
        elif request_user_id and str(request_user_id) == host_user_id:
            has_permission = True

    if not has_permission:
        return JsonResponse(
            {"ok": False, "error": "태스크를 저장할 권한이 없습니다."},
            status=403,
        )

    try:
        payload = json.loads(request.body.decode("utf-8"))
    except json.JSONDecodeError:
        return JsonResponse({"ok": False, "error": "Invalid JSON"}, status=400)

    tasks_payload = payload.get("tasks")
    if tasks_payload is None or not isinstance(tasks_payload, list):
        return JsonResponse({"ok": False, "error": "tasks 형식이 올바르지 않습니다."}, status=400)

    new_tasks = []
    for item in tasks_payload:
        if not isinstance(item, dict):
            continue
        who = (item.get("who") or item.get("assignee") or "").strip()
        assignee_id_payload = item.get("assignee_id")
        what = (item.get("what") or item.get("description") or "").strip()
        when = (item.get("when") or item.get("due") or item.get("due_text") or "").strip()

        # 빈 행은 저장하지 않는다.
        if not (who or what or when):
            continue

        parsed_due = _parse_due_date(when) if when else None
        assignee_obj = None
        if assignee_id_payload:
            try:
                assignee_obj = User.objects.filter(user_id=assignee_id_payload).first()
            except Exception:
                assignee_obj = None
        if not assignee_obj and who:
            try:
                # "이름 (부서)" 형태라면 이름만 떼어 검색
                who_for_lookup = who.split("(", 1)[0].strip()
                # 호스트 이름이 명시되었으면 우선 호스트로 매칭
                if meeting.host and who_for_lookup == meeting.host.name:
                    assignee_obj = meeting.host
                if not assignee_obj:
                    assignee_obj = User.objects.filter(name=who_for_lookup).first()
            except Exception:
                assignee_obj = None

        content_payload = {"description": what}
        if who:
            content_payload["assignee"] = who
            if assignee_obj and getattr(assignee_obj, "user_id", None):
                content_payload["assignee_id"] = assignee_obj.user_id
        if when and when.strip().lower() not in ("*", "null", "none", ""):
            content_payload["due"] = when
            if parsed_due:
                content_payload["due_date"] = parsed_due.isoformat()

        new_tasks.append(
            Task(
                meeting=meeting,
                task_content=json.dumps(content_payload, ensure_ascii=False),
                assignee=assignee_obj,
                due_date=parsed_due,
            )
        )

    with transaction.atomic():
        meeting.tasks.all().delete()
        if new_tasks:
            Task.objects.bulk_create(new_tasks)

    return JsonResponse({"ok": True})


def parse_minutes_sections(html: str) -> Dict[str, str]:
    """
    meeting.meeting_notes 에 저장된 HTML에서 data-minutes-section 별 텍스트 추출
    """
    if not html:
        return {}

    pattern = re.compile(
        r'<div[^>]*data-minutes-section="(?P<key>[^"]+)"[^>]*>(?P<body>.*?)</div>',
        re.DOTALL | re.IGNORECASE,
    )

    sections: Dict[str, str] = {}
    for match in pattern.finditer(html):
        key = match.group("key")
        body_html = match.group("body")
        text = strip_tags(body_html).replace("&nbsp;", " ")
        lines = [ln.rstrip() for ln in text.splitlines()]
        text_clean = "\n".join(ln for ln in lines if ln.strip())
        sections[key] = text_clean

    return sections

def extract_major_agenda(html: str) -> str:
    """
    회의록 HTML 안에서 '주요안건' 행 아래 minutes-editbox 내용을 우선 추출
    """
    if not html:
        return ""

    pattern = re.compile(
        r'<th[^>]*>\s*주요안건\s*</th>.*?<div[^>]*class="[^"]*minutes-editbox[^"]*"[^>]*>(?P<body>.*?)</div>',
        re.DOTALL | re.IGNORECASE,
    )
    match = pattern.search(html)
    if not match:
        return ""

    text = strip_tags(match.group("body")).replace("&nbsp;", " ").strip()
    return text


def minutes_download(request, meeting_id, fmt):
    meeting = get_object_or_404(Meeting, pk=meeting_id)

    attendees_qs = (
        meeting.attendees
               .select_related("user", "user__dept")
               .all()
    )
    attendees_count = attendees_qs.count()

    sections = parse_minutes_sections(meeting.meeting_notes or "")

    raw_contents = sections.get("contents", "")
    raw_results = sections.get("results", "")
    raw_todos = sections.get("todos", "")
    raw_base = sections.get("base", "")

    def clean_section(text: str, header_keywords) -> str:
        if not text:
            return ""
        lines = [ln.rstrip() for ln in text.splitlines()]
        while lines and not lines[0].strip():
            lines.pop(0)

        if lines:
            first = lines[0].replace(" ", "")
            if all(keyword in first for keyword in header_keywords):
                lines.pop(0)

        return "\n".join(lines).strip()

    contents_text = clean_section(raw_contents, ["회의", "내용"])
    results_text = clean_section(raw_results, ["회의", "결과"])
    todos_text = clean_section(raw_todos, ["해야", "할", "일"])
    base_text = clean_section(raw_base, ["기본", "정보"])



    has_base = "base" in sections
    has_contents = "contents" in sections
    has_results = "results" in sections
    has_todos = "todos" in sections
    has_attendees_section = "attendees" in sections

    main_agenda = extract_major_agenda(meeting.meeting_notes or "")
    if not main_agenda and base_text:
        for ln in base_text.splitlines():
            if ln.strip():
                main_agenda = ln.strip()
                break
    if not main_agenda and contents_text:
        for ln in contents_text.splitlines():
            if ln.strip():
                main_agenda = ln.strip()
                break

    agenda = meeting.title or ""
    meeting_dt_str = (
        meeting.meet_date_time.strftime("%Y.%m.%d %H:%M")
        if meeting.meet_date_time
        else ""
    )
    place = meeting.place or ""
    if hasattr(meeting, "host") and meeting.host:
        host_name = meeting.host.name
    else:
        host_name = getattr(meeting, "responsible_name", "")

    if meeting.meeting_notes:
        has_full_minutes = True
        raw_html = (
            meeting.meeting_notes
            .replace("<br>", "\n")
            .replace("<br/>", "\n")
            .replace("<br />", "\n")
        )
        text_body = strip_tags(raw_html).replace("&nbsp;", " ")
    else:
        has_full_minutes = False
        if meeting.summary:
            text_body = meeting.summary
        elif meeting.transcript:
            text_body = meeting.transcript
        else:
            text_body = "회의 내용이 아직 등록되지 않았습니다."

    if fmt == "pdf":
        buffer = BytesIO()
        p = canvas.Canvas(buffer, pagesize=A4)

        width, height = A4
        x_margin = 50
        top_margin = height - 50
        bottom_margin = 40
        line_height = 14

        font = KOREAN_FONT_NAME

        def wrap_line_by_width(line: str, max_width: float, font_size: int = 10):
            """
            주어진 폭 안에 단어 단위로 줄바꿈. 너무 긴 단어는 폭에 맞게 강제 분리.
            """
            words = line.split()
            if not words:
                return [""]
            lines_local = []
            current = words[0]
            for word in words[1:]:
                candidate = f"{current} {word}"
                if pdfmetrics.stringWidth(candidate, font, font_size) <= max_width:
                    current = candidate
                else:
                    lines_local.append(current)
                    current = word
            if pdfmetrics.stringWidth(current, font, font_size) > max_width:
                tmp = current
                while pdfmetrics.stringWidth(tmp, font, font_size) > max_width:
                    cut = len(tmp)
                    while cut > 0 and pdfmetrics.stringWidth(tmp[:cut], font, font_size) > max_width:
                        cut -= 1
                    if cut <= 0:
                        break
                    lines_local.append(tmp[:cut])
                    tmp = tmp[cut:]
                if tmp:
                    lines_local.append(tmp)
            else:
                lines_local.append(current)
            return lines_local

        def draw_page_border(top_y, bottom_y):
            p.line(x_left, top_y, x_left, bottom_y)
            p.line(x_right, top_y, x_right, bottom_y)

        def start_new_page():
            nonlocal current_top, table_top_y
            p.setFont(font, 11)
            current_top = top_margin
            table_top_y = current_top

        def finalize_current_page(outer_bottom_value):
            draw_page_border(table_top_y, outer_bottom_value)

        y = top_margin
        p.setFont(font, 18)
        p.drawString(x_margin, y, meeting.title or "회의록")
        y -= 40

        p.setFont(font, 11)
        x_left = x_margin
        x_right = width - x_margin

        header_row_h = 24
        label_row_h = 24
        contents_h = 200
        results_h = 200
        todos_h = 150
        attend_header_h = 24
        attend_row_h = 24

        rows_per_side = max(4, math.ceil(attendees_count / 2))
        attend_rows_h = rows_per_side * attend_row_h

        table_top_y = y
        current_top = table_top_y

        center_x = (x_left + x_right) / 2.0

        if has_base:
            label_w = 80
            right_w = 120
            x_label = x_left + label_w
            x_right_block = x_right - right_w

            header_top = current_top

            for i in range(5):
                y_line = header_top - header_row_h * i
                p.line(x_left, y_line, x_right, y_line)

            p.line(x_left, header_top, x_left, header_top - 4 * header_row_h)
            p.line(x_right, header_top, x_right, header_top - 4 * header_row_h)

            p.line(x_label, header_top, x_label, header_top - header_row_h)
            p.line(x_label, header_top - header_row_h, x_label, header_top - 2 * header_row_h)
            p.line(x_label, header_top - 2 * header_row_h, x_label, header_top - 3 * header_row_h)
            p.line(x_label, header_top - 3 * header_row_h, x_label, header_top - 4 * header_row_h)

            p.line(x_right_block, header_top - header_row_h, x_right_block, header_top - 3 * header_row_h)

            def header_text_y(row_idx: int) -> float:
                return header_top - header_row_h * row_idx - header_row_h + 8

            y_row1 = header_text_y(0)
            p.drawString(x_left + 5, y_row1, "제 목")
            p.drawString(x_left + 5 + 80, y_row1, agenda)

            y_row2 = header_text_y(1)
            p.drawString(x_left + 5, y_row2, "일 시")
            p.drawString(x_left + 5 + 80, y_row2, meeting_dt_str)
            host_label = "주최자명"
            if host_name:
                host_label += f" {host_name}"
            p.drawString(x_right_block + 5, y_row2, host_label)

            y_row3 = header_text_y(2)
            p.drawString(x_left + 5, y_row3, "장 소")
            p.drawString(x_left + 5 + 80, y_row3, place)
            p.drawString(x_right_block + 5, y_row3, f"참석인원  {attendees_count}")

            y_row4 = header_text_y(3)
            p.drawString(x_left + 5, y_row4, "주요안건")
            if main_agenda:
                agenda_box_width = x_right - (x_left + 80) - 10
                agenda_lines = wrap_line_by_width(main_agenda, agenda_box_width, font_size=11)
                agenda_y = y_row4
                for line in agenda_lines:
                    p.drawString(x_left + 5 + 80, agenda_y, line)
                    agenda_y -= line_height

            current_top = header_top - 4 * header_row_h
        else:
            current_top = y

        if has_contents:
            contents_title_top = current_top
            contents_title_bottom = contents_title_top - label_row_h
            p.line(x_left, contents_title_bottom, x_right, contents_title_bottom)
            p.drawCentredString(center_x, contents_title_bottom + 8, "회의 내용")

            contents_box_top = contents_title_bottom
            contents_box_bottom = contents_box_top - contents_h
            p.line(x_left, contents_box_bottom, x_right, contents_box_bottom)

            current_top = contents_box_bottom
        else:
            contents_box_top = contents_box_bottom = None

        if has_results:
            results_title_top = current_top
            results_title_bottom = results_title_top - label_row_h
            p.line(x_left, results_title_bottom, x_right, results_title_bottom)
            p.drawCentredString(center_x, results_title_bottom + 8, "회의 결과")

            results_box_top = results_title_bottom
            results_box_bottom = results_box_top - results_h
            p.line(x_left, results_box_bottom, x_right, results_box_bottom)

            current_top = results_box_bottom
        else:
            results_box_top = results_box_bottom = None

        if has_todos:
            todos_title_top = current_top
            todos_title_bottom = todos_title_top - label_row_h
            p.line(x_left, todos_title_bottom, x_right, todos_title_bottom)
            p.drawCentredString(center_x, todos_title_bottom + 8, "해야 할 일")

            todos_box_top = todos_title_bottom
            todos_box_bottom = todos_box_top - todos_h
            p.line(x_left, todos_box_bottom, x_right, todos_box_bottom)

            current_top = todos_box_bottom
        else:
            todos_box_top = todos_box_bottom = None

        p.setFont(font, 10)

        def draw_multiline_in_box(text, x_left_box, x_right_box, top_y, bottom_y):
            usable_width = (x_right_box - x_left_box) - 10  # 좌우 여백 5씩 확보
            y_pos = top_y - 14
            for raw_line in (text or "").splitlines():
                wrapped = wrap_line_by_width(raw_line, usable_width)
                for line in wrapped:
                    if not line and len(wrapped) == 1:
                        continue
                    if y_pos < bottom_y + line_height:
                        return
                    p.drawString(x_left_box + 5, y_pos, line)
                    y_pos -= line_height

        if has_contents and contents_box_top is not None:
            draw_multiline_in_box(
                contents_text,
                x_left,
                x_right,
                contents_box_top,
                contents_box_bottom,
            )

        if has_results and results_box_top is not None:
            draw_multiline_in_box(
                results_text,
                x_left,
                x_right,
                results_box_top,
                results_box_bottom,
            )

        if has_todos and todos_box_top is not None:
            draw_multiline_in_box(
                todos_text,
                x_left,
                x_right,
                todos_box_top,
                todos_box_bottom,
            )

        p.setFont(font, 11)

        if has_attendees_section or not sections:
            attend_label_top = current_top
            attend_label_bottom = attend_label_top - label_row_h
            attend_block_height = label_row_h + attend_header_h + attend_rows_h

            if attend_label_top - attend_block_height < bottom_margin:
                outer_bottom = current_top
                finalize_current_page(outer_bottom)
                p.showPage()
                start_new_page()
                attend_label_top = current_top
                attend_label_bottom = attend_label_top - label_row_h

            p.line(x_left, attend_label_bottom, x_right, attend_label_bottom)
            p.drawCentredString(center_x, attend_label_bottom + 8, "참석자")

            attend_header_top = attend_label_bottom
            attend_header_bottom = attend_header_top - attend_header_h
            p.line(x_left, attend_header_bottom, x_right, attend_header_bottom)

            col_w = (x_right - x_left) / 6.0
            x_cols = [x_left + col_w * i for i in range(7)]

            attend_table_bottom = attend_header_bottom - attend_rows_h
            for xv in x_cols:
                p.line(xv, attend_header_top, xv, attend_table_bottom)

            header_y = attend_header_bottom + 5
            p.drawString(x_cols[0] + 5, header_y, "소 속")
            p.drawString(x_cols[1] + 5, header_y, "성 명")
            p.drawString(x_cols[2] + 5, header_y, "서 명")
            p.drawString(x_cols[3] + 5, header_y, "소 속")
            p.drawString(x_cols[4] + 5, header_y, "성 명")
            p.drawString(x_cols[5] + 5, header_y, "서 명")

            row_top = attend_header_bottom
            for _ in range(rows_per_side):
                row_top -= attend_row_h
                p.line(x_left, row_top, x_right, row_top)

            outer_bottom = attend_table_bottom
        else:
            outer_bottom = current_top

        finalize_current_page(outer_bottom)

        if has_attendees_section or not sections:
            attendees_list = list(attendees_qs)
            rows_per_side = max(4, math.ceil(len(attendees_list) / 2))
            col_w = (x_right - x_left) / 6.0
            x_cols = [x_left + col_w * i for i in range(7)]
            for row_idx in range(rows_per_side):
                row_text_y = attend_header_bottom - attend_row_h * row_idx - attend_row_h + 5

                left_idx = row_idx
                if left_idx < len(attendees_list):
                    att = attendees_list[left_idx]
                    dept = att.user.dept.dept_name if att.user.dept else ""
                    name = att.user.name
                    p.drawString(x_cols[0] + 5, row_text_y, dept)
                    p.drawString(x_cols[1] + 5, row_text_y, name)
                p.drawString(x_cols[2] + 5, row_text_y, "(인)")

                right_idx = row_idx + rows_per_side
                if right_idx < len(attendees_list):
                    att = attendees_list[right_idx]
                    dept = att.user.dept.dept_name if att.user.dept else ""
                    name = att.user.name
                    p.drawString(x_cols[3] + 5, row_text_y, dept)
                    p.drawString(x_cols[4] + 5, row_text_y, name)
                p.drawString(x_cols[5] + 5, row_text_y, "(인)")

        p.showPage()
        p.save()

        pdf_value = buffer.getvalue()
        buffer.close()

        filename = f"meeting_{meeting_id}_minutes.pdf"
        response = HttpResponse(pdf_value, content_type="application/pdf")
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response

    if fmt == "docx":
        doc = Document()
        doc.add_heading(meeting.title or "회의록", level=1)

        info_rows = [
            ("제목", agenda or "-"),
            ("일시", meeting_dt_str or "-"),
            ("장소", place or "-"),
            ("주최자", host_name or "-"),
            ("주요안건", main_agenda or "-"),
            ("참석자 수", str(attendees_count)),
        ]
        info_table = doc.add_table(rows=len(info_rows), cols=2)
        info_table.style = "Table Grid"
        for idx, (label, value) in enumerate(info_rows):
            cells = info_table.rows[idx].cells
            cells[0].text = label
            cells[1].text = value

        def add_doc_section(title: str, text: str):
            cleaned = (text or "").strip()
            if not cleaned:
                return
            doc.add_paragraph("")
            doc.add_heading(title, level=2)
            for ln in cleaned.splitlines():
                doc.add_paragraph(ln)

        add_doc_section("회의 내용", contents_text)
        add_doc_section("회의 결과", results_text)
        add_doc_section("해야 할 일", todos_text)

        doc.add_paragraph("")
        doc.add_heading("참석자", level=2)
        attendees_list = list(attendees_qs)
        if attendees_list:
            rows = max(1, math.ceil(len(attendees_list) / 2))
            table = doc.add_table(rows=rows + 1, cols=4)
            table.style = "Table Grid"
            header_cells = table.rows[0].cells
            header_cells[0].text = "소 속"
            header_cells[1].text = "성 명"
            header_cells[2].text = "소 속"
            header_cells[3].text = "성 명"

            for row_idx in range(rows):
                row_cells = table.rows[row_idx + 1].cells
                left_idx = row_idx
                if left_idx < len(attendees_list):
                    att = attendees_list[left_idx]
                    dept = att.user.dept.dept_name if att.user.dept else ""
                    name = att.user.name
                    row_cells[0].text = dept
                    row_cells[1].text = name
                right_idx = row_idx + rows
                if right_idx < len(attendees_list):
                    att = attendees_list[right_idx]
                    dept = att.user.dept.dept_name if att.user.dept else ""
                    name = att.user.name
                    row_cells[2].text = dept
                    row_cells[3].text = name
        else:
            doc.add_paragraph("참석자 정보가 없습니다.")

        if not meeting.meeting_notes and text_body:
            doc.add_paragraph("")
            doc.add_heading("회의 내용", level=2)
            for line in text_body.splitlines():
                doc.add_paragraph(line)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)

        filename = f"meeting_{meeting_id}_minutes.docx"
        response = HttpResponse(
            buf.getvalue(),
            content_type=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response

    return HttpResponse("invalid format", status=400)
